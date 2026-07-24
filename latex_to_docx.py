#!/usr/bin/env python3
r"""
LaTeX to Word (.docx) Converter
================================
Converts a LaTeX .tex file to a Word document preserving:
  - Title page  (title image, title, author, date)
  - Header / footer (logo + project name | page number)
  - Sections, subsections, subsubsections, paragraphs
  - Custom \newcommand macros and \setcounter values
  - Paragraphs with bold / italic / inline math / hyperlinks
  - Itemize and enumerate lists (including custom labels)
  - longtable and tabularx/tabular environments
  - Display math  \[ ... \]  rendered as centred italic text
  - Page breaks (\newpage, \clearpage)
  - Table of contents placeholder

Usage:
    python latex_to_docx.py [input.tex] [output.docx]
    Defaults: linea_base_en.tex  ->  linea_base_en.docx
"""

import re
import sys
import subprocess
import shutil
import tempfile
import unicodedata
from pathlib import Path
from copy import deepcopy

from dataclasses import dataclass, field
from docx import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
from lxml import etree


# ─────────────────────────────────────────────────────────────────────────────
# Configuration
# ─────────────────────────────────────────────────────────────────────────----

@dataclass
class Config:
    """Runtime configuration populated from the LaTeX preamble and CLI defaults."""
    # Presentation
    font_name: str = "Calibri"
    font_size_normal: int = 11
    font_size_small: int = 9
    font_size_foot: int = 8

    # Page
    page_width_inches: float = 8.27
    page_height_inches: float = 11.69
    margin_top = Cm(2.5)
    margin_bottom = Cm(2.0)
    margin_left = Cm(2.54)
    margin_right = Cm(2.54)

    # Calculated from page + margins
    text_width_inches: float = 6.27
    text_height_inches: float = 9.50

    # Language / captions
    language: str = "es"
    table_name: str = "Table"
    figure_name: str = "Figure"
    chapter_prefix: str = "Capítulo "
    appendix_prefix: str = "Apéndice "
    bib_title_default: str = "Referencias"

    # TOC/LOT/LOF titles and placeholders
    toc_title: str = "TABLE OF CONTENTS"
    lof_title: str = "LIST OF FIGURES"
    lot_title: str = "LIST OF TABLES"
    toc_placeholder: str = "Right-click and select Update Field to generate."
    lof_placeholder: str = "No figure entries found. Right-click and select Update Field to generate."
    lot_placeholder: str = "No table entries found. Right-click and select Update Field to generate."

    # Executables
    pandoc_path: str = "pandoc"
    pdflatex_path: str = "pdflatex"
    pdftoppm_path: str = "pdftoppm"

    # Timeouts (seconds)
    timeout_pandoc: int = 60
    timeout_pdflatex: int = 90
    timeout_pdftoppm: int = 30
    timeout_math: int = 30

    # Images
    image_dpi: float = 96.0
    image_tikz_dpi: int = 300
    image_default_width_inches: float = 4.0
    image_max_width_inches: float = 6.3

    # Logo search candidates (relative to input .tex directory)
    logo_candidates: list[str] = field(default_factory=lambda: [
        "icono.png", "icono.jpg",
        "images/logo.jpg", "images/logo.png",
        "00_figs/logo.jpg",
    ])
    logo_height = Cm(0.8)

    # Colours
    color_black = RGBColor(0x00, 0x00, 0x00)
    color_white = RGBColor(0xFF, 0xFF, 0xFF)
    color_link = RGBColor(0x00, 0x56, 0xB3)
    color_gray = RGBColor(0x60, 0x60, 0x60)
    hex_header = "1F497D"
    hex_category = "DDEEFF"
    hex_alt = "F5F9FF"

    # Fallback natural image size when Pillow cannot open the file
    fallback_image_width: float = 4.0
    fallback_image_height: float = 3.0


# Global configuration instance; populated in main() from the LaTeX preamble.
CONFIG: Config = Config()

# ─────────────────────────────────────────────────────────────────────────────
# Colours (kept as module-level constants for now)
# ─────────────────────────────────────────────────────────────────────────────
C_BLACK   = RGBColor(0x00, 0x00, 0x00)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_LINK    = RGBColor(0x00, 0x56, 0xB3)
C_GRAY    = RGBColor(0x60, 0x60, 0x60)

HEX_HDR   = "1F497D"   # dark blue for table header rows
HEX_CAT   = "DDEEFF"   # light blue for category rows inside tables
HEX_ALT   = "F5F9FF"   # very light blue alternating row (unused but available)

FONT      = "Calibri"
PT_NORM   = 11
PT_SMALL  =  9

# Carpeta temporal para archivos intermedios (tablas, etc.)
TEMP_DIR: Path = Path(".")  # Se establece en main()
PT_FOOT   =  8

PANDOC_PATH: str = "pandoc"  # Se establece en main()
MATH_NS: str = "http://schemas.openxmlformats.org/officeDocument/2006/math"
OMML_CACHE: dict[str, object] = {}

# Ancho de texto para este documento (A4, margin=1in → 8.27 - 2 = 6.27in)
TEXT_WIDTH_INCHES  = 6.27
# Alto aproximado del cuerpo de texto (A4, margin ~1in arriba/abajo → 11.69 - 2 ≈ 9.5in)
TEXT_HEIGHT_INCHES = 9.50

# ─────────────────────────────────────────────────────────────────────────────
# XML helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color="BBBBBB", sz="4"):
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    tcBdr = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:color"), color)
        tcBdr.append(el)
    tcPr.append(tcBdr)


def _add_page_number_field(run):
    """Insert a PAGE field into a run (with proper begin/separate/end structure)."""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    # Placeholder display text (updated by Word on open)
    t = OxmlElement("w:t")
    t.text = "1"
    run._r.append(t)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _add_tc_field(paragraph, text: str, identifier: str):
    """Add a hidden TC (Table of Contents Entry) field to a paragraph."""
    safe_text = text.replace('"', "'")

    run1 = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run1._r.append(fld_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = f' TC "{safe_text}" \\f {identifier} '
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_sep)

    run4 = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run4._r.append(fld_end)


def _safe_page_break(doc: Document):
    """Add a page break only if the document does not already end with one.
    Also strips trailing empty paragraphs to avoid blank pages."""
    # Remove trailing completely empty paragraphs (no text, no runs, no math)
    while doc.paragraphs:
        last_p = doc.paragraphs[-1]
        has_content = bool(last_p.text.strip()) or len(last_p.runs) > 0
        if not has_content:
            # Equations are added as OMML children, not runs; keep those paragraphs.
            for child in last_p._p:
                tag = child.tag
                if tag == f"{{{MATH_NS}}}oMath" or tag == f"{{{MATH_NS}}}oMathPara":
                    has_content = True
                    break
        if not has_content:
            last_p._element.getparent().remove(last_p._element)
        else:
            break

    # Check if the last paragraph already ends with a page break
    if doc.paragraphs:
        last_p = doc.paragraphs[-1]
        for run in last_p.runs:
            for br in run._r.findall(qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    return  # Already a page break, skip

    doc.add_page_break()


def _remove_empty_paragraphs_after_tables(doc: Document):
    """Leave one compact empty paragraph after each table.

    Pandoc usually emits one or more empty paragraphs right after a table.
    Removing all of them makes the text look glued to the table; keeping the
    first one with a small space_after gives a predictable, comfortable gap.
    """
    body = doc.element.body
    removed = 0
    compacted = 0
    # Recorrer de atrás hacia adelante para que eliminar párrafos no cambie
    # el hermano siguiente de las tablas anteriores.
    for tbl in reversed(list(body.findall(qn('w:tbl')))):
        empties = []
        nxt = tbl.getnext()
        while nxt is not None and nxt.tag == qn('w:p'):
            text = ''.join(nxt.itertext()).strip()
            if not text:
                empties.append(nxt)
                nxt = nxt.getnext()
            else:
                break

        if not empties:
            continue

        # Keep the first empty paragraph, delete any extras.
        for p in empties[1:]:
            p.getparent().remove(p)
            removed += 1

        # Compact the remaining separator paragraph but keep a small gap.
        sep = Paragraph(empties[0], doc)
        sep.paragraph_format.space_after = Pt(6)
        sep.paragraph_format.space_before = Pt(0)
        sep.paragraph_format.line_spacing = 1.0
        compacted += 1

    if removed or compacted:
        print(f"  [INFO] Ajustados {compacted} espacios tras tablas, eliminados {removed} párrafos vacíos extra")


def _insert_toc_field(doc: Document, instr: str, title: str = "", placeholder: str = ""):
    """Insert a Word TOC/LOT/LOF field."""
    from docx.oxml import parse_xml

    # NOTE: We intentionally do NOT set w:updateFields to true because that
    # triggers a Word dialog on every open asking whether to update fields.
    # The user can update once with Ctrl+A → F9, save, and the document
    # will open cleanly afterwards.

    if title:
        p = doc.add_paragraph()
        run = _run(p, title, bold=True, size=13)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.space_before = Pt(12)
    toc_para.paragraph_format.space_after = Pt(12)

    toc_para._p.append(parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:fldChar w:fldCharType="begin"/></w:r>'
    ))
    toc_para._p.append(parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:instrText xml:space="preserve"> {instr} </w:instrText></w:r>'
    ))
    toc_para._p.append(parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:fldChar w:fldCharType="separate"/></w:r>'
    ))
    ph_text = placeholder or "Right-click and select Update Field to generate."
    toc_para._p.append(parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:rPr><w:color w:val="808080"/><w:i/></w:rPr>'
        f'<w:t>{ph_text}</w:t></w:r>'
    ))
    toc_para._p.append(parse_xml(
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:fldChar w:fldCharType="end"/></w:r>'
    ))


def add_toc(doc: Document, title: str | None = None, config: Config | None = None):
    """Insert a Table of Contents field (auto-updates on open)."""
    if config is None:
        config = CONFIG
    _insert_toc_field(
        doc, r'TOC \o "1-3" \h \z',
        title if title is not None else config.toc_title,
        config.toc_placeholder
    )


def add_lof(doc: Document, title: str | None = None, config: Config | None = None):
    """Insert a List of Figures field (auto-updates on open)."""
    if config is None:
        config = CONFIG
    _insert_toc_field(
        doc, r'TOC \c "Figure" \h \z',
        title if title is not None else config.lof_title,
        config.lof_placeholder
    )


def add_lot(doc: Document, title: str | None = None, config: Config | None = None):
    """Insert a List of Tables field (auto-updates on open)."""
    if config is None:
        config = CONFIG
    _insert_toc_field(
        doc, r'TOC \c "Table" \h \z',
        title if title is not None else config.lot_title,
        config.lot_placeholder
    )


def add_table_caption(doc: Document, caption_text: str, table_num: int = 0,
                      config: Config | None = None):
    """
    Add a table caption with a hidden TC field for Word's List of Tables.
    """
    if config is None:
        config = CONFIG

    clean_caption = strip_fmt(caption_text).replace('"', "'")
    prefix = config.table_name
    full_caption = f"{prefix} {table_num}: {clean_caption}" if table_num > 0 else f"{prefix}: {clean_caption}"

    # Simple paragraph - no special style
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Visible caption
    r = p.add_run(full_caption)
    r.bold = True
    r.font.size = Pt(config.font_size_small)
    r.font.name = config.font_name
    r.font.color.rgb = C_BLACK

    # Hidden TC field for List of Tables
    tc_text = f"{prefix} {table_num}\t{clean_caption}" if table_num > 0 else f"{prefix}\t{clean_caption}"
    _add_tc_field(p, tc_text, "T")


def add_list_of_tables(doc: Document, title: str = "LIST OF TABLES"):
    r"""
    Insert a List of Tables field using TOA (Table of Authorities) approach.
    This creates a proper List of Tables that Word can update.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Title
    p = doc.add_paragraph()
    run = _run(p, title, bold=True, size=13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Create paragraph for the LOT field
    lot_para = doc.add_paragraph()
    lot_para.paragraph_format.space_before = Pt(12)
    lot_para.paragraph_format.space_after = Pt(12)
    
    # Field: TOA \h \c "Table"  (Table of Authorities for Tables)
    r = lot_para.add_run()
    
    # fldChar begin
    e1 = OxmlElement('w:fldChar')
    e1.set(qn('w:fldCharType'), 'begin')
    r._r.append(e1)
    
    # instrText - TOA (Table of Authorities) with \c category
    e2 = OxmlElement('w:instrText')
    e2.set(qn('xml:space'), 'preserve')
    e2.text = ' TOA \\h \\c "1" '
    r._r.append(e2)
    
    # fldChar separate
    e3 = OxmlElement('w:fldChar')
    e3.set(qn('w:fldCharType'), 'separate')
    r._r.append(e3)
    
    # Placeholder text
    e4 = OxmlElement('w:r')
    e4_pr = OxmlElement('w:rPr')
    e4_color = OxmlElement('w:color')
    e4_color.set(qn('w:val'), '808080')
    e4_pr.append(e4_color)
    e4_i = OxmlElement('w:i')
    e4_pr.append(e4_i)
    e4.append(e4_pr)
    e4_t = OxmlElement('w:t')
    e4_t.text = '[Press Ctrl+A, F9 to update List of Tables]'
    e4.append(e4_t)
    lot_para._p.append(e4)
    
    # fldChar end
    e5 = OxmlElement('w:fldChar')
    e5.set(qn('w:fldCharType'), 'end')
    r._r.append(e5)

    # Hint paragraph
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Text helpers
# ─────────────────────────────────────────────────────────────────────────────

def _run(para, text: str, bold=False, italic=False,
         size=None, color=None, underline=False):
    if not text:
        return None
    r = para.add_run(text)
    r.font.name   = FONT
    r.font.size   = Pt(size or PT_NORM)
    r.bold        = bold
    r.italic      = italic
    r.underline   = underline
    if color:
        r.font.color.rgb = color
    return r


# ─────────────────────────────────────────────────────────────────────────────
# LaTeX text cleaning
# ─────────────────────────────────────────────────────────────────────────────

# Populated during preamble parsing
MACROS: dict[str, str] = {}


def _resolve_macros(text: str) -> str:
    """Replace custom commands with their resolved text.
    Sort by key length descending so \\PrestadorII replaces before \\PrestadorI.
    """
    for k in sorted(MACROS.keys(), key=len, reverse=True):
        text = text.replace(k, MACROS[k])
    return text


def _clean(text: str) -> str:
    """
    Convert LaTeX special chars/sequences to Unicode.
    Does NOT strip formatting commands — use strip_fmt() for plain text.
    """
    text = _resolve_macros(text)
    # Note: font-size switches (\large, \normalsize, etc.) and font switches
    # (\bfseries, \normalfont) are now handled by parse_inline so they can
    # affect formatting. They are no longer stripped here.
    # Remove \centering to avoid \c being treated as cedilla during accent processing.
    text = re.sub(r"\\centering\b", " ", text)
    # Convert LaTeX accent commands to Unicode (\'{a}, \'a, \~{n}, \^u, etc.)
    # MUST happen before replacing standalone ~ with non-breaking space
    def _accent_repl(m):
        accent_cmd = m.group(1)
        letter = m.group(2)
        combining = {
            "'": "\u0301", "`": "\u0300", "^": "\u0302",
            '"': "\u0308", "~": "\u0303", "c": "\u0327",
            "v": "\u030c",  # caron
        }.get(accent_cmd, "")
        return letter + combining if combining else m.group(0)

    text = re.sub(r"\\(['`^\"~cv])\{?([a-zA-Z])\}?", _accent_repl, text)
    text = unicodedata.normalize("NFC", text)
    # Now safe to replace standalone ~ (non-breaking space), ---, --, etc.
    text = text.replace("~",   "\u00a0")          # non-breaking space
    text = text.replace("---", "\u2014")           # em dash
    text = text.replace("--",  "\u2013")           # en dash
    text = text.replace("\\%", "%")
    text = text.replace("\\$", "$")
    text = text.replace("\\&", "&")
    text = text.replace("\\_", "_")
    text = text.replace("\\#", "#")
    text = text.replace("\\,", "\u202f")           # narrow no-break space
    text = re.sub(r"\\hspace\*?\{[^}]*\}", "\t", text)  # replace \hspace{...} with tab
    text = re.sub(r"\\vspace\*?\{[^}]*\}", "\t", text)  # replace \vspace{...} with tab
    # Collapse whitespace surrounding any tab into a single tab
    text = re.sub(r"[ \t]*\t[ \t]*", "\t", text)
    text = re.sub(r"\\[,;!: ]\s*", " ", text)
    text = re.sub(r"\\ ", " ", text)
    return text


def strip_fmt(text: str) -> str:
    """Return plain text: resolve macros + remove all LaTeX markup."""
    text = _clean(text)
    # Strip inline math $...$ — replace \cdot etc. with unicode, then strip $
    text = re.sub(r"\$\\cdot\$",             "·",   text)
    text = re.sub(r"\$\\approx\$",           "≈",   text)
    text = re.sub(r"\$\\times\$",            "×",   text)
    text = re.sub(r"\$([^$]*)\$",            r"\1", text)   # remaining $math$
    text = re.sub(r"\\textbf\{([^}]*)\}",    r"\1", text)
    text = re.sub(r"\\textit\{([^}]*)\}",    r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}",      r"\1", text)
    text = re.sub(r"\\textsc\{([^}]*)\}",    r"\1", text)  # \textsc{...}
    text = re.sub(r"\\small\b",              "",    text)
    text = re.sub(r"\\large\b",              "",    text)
    text = re.sub(r"\\Large\b",              "",    text)
    text = re.sub(r"\\normalsize\b",         "",    text)
    text = re.sub(r"\\footnotesize\b",       "",    text)
    text = re.sub(r"\\scriptsize\b",         "",    text)
    text = re.sub(r"\\tiny\b",               "",    text)
    text = re.sub(r"\\huge\b",               "",    text)
    text = re.sub(r"\\Huge\b",               "",    text)
    text = re.sub(r"\\centering\b",          "",    text)
    text = re.sub(r"\\textbullet\b",         "•",   text)
    text = re.sub(r"\\href\{[^}]*\}\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\url\{([^}]*)\}",       r"\1", text)
    text = re.sub(r"\\ref\{[^}]*\}",         "",    text)
    text = re.sub(r"\\label\{[^}]*\}",       "",    text)
    text = re.sub(r"\\newline\b",            " ",   text)
    text = re.sub(r"\\hspace\*?\{[^}]*\}",  " ",   text)  # replace \hspace{...} with space
    text = re.sub(r"\\vspace\*?\{[^}]*\}",  " ",   text)  # replace \vspace{...} with space
    text = re.sub(r"\\rule\{[^}]*\}\{[^}]*\}", "", text)  # remove \rule{...}{...}
    # Line breaks with optional argument: \\[\bigskipamount] → space
    text = re.sub(r"\\\\(?:\[[^\]]*\])?", " ", text)
    text = re.sub(r"\\includegraphics(?:\[[^\]]*\])?\{[^}]*\}", "", text)
    text = re.sub(r"\\[a-zA-Z]+\*?\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?",        "",    text)
    text = text.replace("{", "").replace("}", "")
    # Normalize internal whitespace (multi-line LaTeX → single line)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


# ─────────────────────────────────────────────────────────────────────────────
# LaTeX math → OMML (native Word equations) via Pandoc
# ─────────────────────────────────────────────────────────────────────────────

MATH_TEMP_COUNTER = [0]

def latex_to_omml_element(latex_code: str, temp_dir: Path, pandoc_path: str):
    r"""
    Convierte código LaTeX matemático a un elemento OMML usando Pandoc.
    `latex_code` debe incluir los delimitadores ($...$, \[...\], o entorno completo).
    Retorna el elemento XML (m:oMath o m:oMathPara) o None.
    Los resultados se cachean para evitar re-conversiones.
    """
    cache_key = latex_code
    if cache_key in OMML_CACHE:
        cached = OMML_CACHE[cache_key]
        if cached is not None:
            # lxml deepcopy is unreliable; serialize + parse for a true independent copy
            return parse_xml(etree.tostring(cached, encoding='utf-8'))
        return None

    MATH_TEMP_COUNTER[0] += 1
    n = MATH_TEMP_COUNTER[0]
    tex_file = temp_dir / f"math_{n:04d}.tex"
    out_file = temp_dir / f"math_{n:04d}.docx"

    # Expand custom macros so Pandoc can parse math that relies on them
    resolved_code = _resolve_macros(latex_code)
    # Pandoc mis-parses the European decimal brace syntax {,} in math mode,
    # splitting the number (e.g. $25{,}8^{\circ}$ becomes 25 , 8°).
    # Replace it with a dot for Pandoc and restore commas in the OMML output.
    resolved_code = resolved_code.replace("{,}", ".")
    # Pandoc also produces empty-base superscripts for unbraced exponents
    # like $m^3$; normalise to $m^{3}$ so the base is preserved.
    resolved_code = re.sub(r"(?<!\{)\^([A-Za-z0-9])(?![A-Za-z0-9])", r"^{\1}", resolved_code)
    tex_content = (
        r"\documentclass{article}" + "\n"
        r"\begin{document}" + "\n"
        + resolved_code + "\n"
        r"\end{document}" + "\n"
    )
    tex_file.write_text(tex_content, encoding="utf-8")

    try:
        subprocess.run(
            [pandoc_path, "-f", "latex", "-t", "docx", str(tex_file), "-o", str(out_file)],
            capture_output=True, text=True, timeout=30
        )
    except Exception:
        return None

    if not out_file.exists():
        return None

    try:
        pandoc_doc = Document(str(out_file))
        for para in pandoc_doc.paragraphs:
            for child in para._p:
                tag = child.tag
                if tag == f"{{{MATH_NS}}}oMath" or tag == f"{{{MATH_NS}}}oMathPara":
                    result = deepcopy(child)
                    # Restore European decimal commas and use the standard
                    # degree sign (U+00B0) instead of the ring operator.
                    for t_el in result.iter(f"{{{MATH_NS}}}t"):
                        if t_el.text:
                            txt = t_el.text
                            txt = txt.replace("\u2218", "\u00b0")
                            txt = re.sub(r"(\d)\.(\d)", r"\1,\2", txt)
                            t_el.text = txt
                    OMML_CACHE[cache_key] = result
                    return result
    except Exception:
        pass

    return None


# ─────────────────────────────────────────────────────────────────────────────
# Inline LaTeX → Word runs
# ─────────────────────────────────────────────────────────────────────────────

# Token patterns (order matters)
_TOK = re.compile(
    r"(\\textbf\{(?:[^{}]|\{[^{}]*\})*\})"           # \textbf{...}
    r"|(\\textit\{(?:[^{}]|\{[^{}]*\})*\})"           # \textit{...}
    r"|(\\emph\{(?:[^{}]|\{[^{}]*\})*\})"             # \emph{...}
    r"|(\\href\{[^}]*\}\{(?:[^{}]|\{[^{}]*\})*\})"   # \href{url}{text}
    r"|(\\url\{[^}]*\})"                               # \url{...}
    r"|(\$[^$]+\$)"                                    # $inline math$
    r"|(\\(?:label|ref|pageref|eqref)\{[^}]*\})"      # \label{...}, \ref{...} — skip
    r"|(\\(?:bfseries|normalfont|itshape|large|Large|LARGE|small|normalsize|footnotesize|scriptsize|tiny|huge|Huge)\b)"  # font switches
    r"|(\\[a-zA-Z]+\*?(?:\{[^}]*\})*)"                # generic \cmd{arg}...
    r"|([^\\${}]+)"                                    # plain text
)


def parse_inline(para, text: str, base_sz=PT_NORM, bold=False, italic=False, size=None):
    r"""Parse inline LaTeX and add styled runs to `para`.

    Supports formatting switches such as \bfseries, \normalfont, \itshape,
    \large, \normalsize, etc. by tracking current bold/italic/size state.
    """
    text = _resolve_macros(text)
    # Protect inline math ($...$) from _clean, which otherwise corrupts
    # commands like \circ, \cdot, etc. before Pandoc can process them.
    math_segments = []
    def _protect_math(m):
        math_segments.append(m.group(0))
        return f"\x00MATH{len(math_segments) - 1}\x00"
    text = re.sub(r"\$[^$]+\$", _protect_math, text)
    text = _clean(text)
    def _restore_math(m):
        return math_segments[int(m.group(1))]
    text = re.sub(r"\x00MATH(\d+)\x00", _restore_math, text)

    cur_bold = bool(bold)
    cur_italic = bool(italic)
    cur_size = size if size is not None else base_sz

    # LaTeX size switches mapped to approximate Word point sizes
    size_map = {
        "tiny": 6,
        "scriptsize": 8,
        "footnotesize": 9,
        "small": 10,
        "normalsize": PT_NORM,
        "large": 14,
        "Large": 16,
        "LARGE": 20,
        "huge": 24,
        "Huge": 28,
    }

    for m in _TOK.finditer(text):
        g0 = m.group(0)
        if m.group(1):   # \textbf
            inner_m = re.search(r"\\textbf\{((?:[^{}]|\{[^{}]*\})*)\}", g0)
            if inner_m:
                n_runs = len(para.runs)
                parse_inline(para, inner_m.group(1), base_sz=base_sz,
                             bold=cur_bold, italic=cur_italic, size=cur_size)
                for run in para.runs[n_runs:]:
                    run.bold = True
        elif m.group(2): # \textit
            inner_m = re.search(r"\\textit\{((?:[^{}]|\{[^{}]*\})*)\}", g0)
            if inner_m:
                n_runs = len(para.runs)
                parse_inline(para, inner_m.group(1), base_sz=base_sz,
                             bold=cur_bold, italic=cur_italic, size=cur_size)
                for run in para.runs[n_runs:]:
                    run.italic = True
        elif m.group(3): # \emph
            inner_m = re.search(r"\\emph\{((?:[^{}]|\{[^{}]*\})*)\}", g0)
            if inner_m:
                n_runs = len(para.runs)
                parse_inline(para, inner_m.group(1), base_sz=base_sz,
                             bold=cur_bold, italic=cur_italic, size=cur_size)
                for run in para.runs[n_runs:]:
                    run.italic = True
        elif m.group(4): # \href{url}{text}
            link_txt = re.sub(r"\\href\{[^}]*\}\{((?:[^{}]|\{[^{}]*\})*)\}", r"\1", g0)
            _run(para, _clean(strip_fmt(link_txt)), color=C_LINK, underline=True,
                 size=cur_size, bold=cur_bold, italic=cur_italic)
        elif m.group(5): # \url{...}
            url = re.sub(r"\\url\{([^}]*)\}", r"\1", g0)
            _run(para, url, color=C_LINK, underline=True,
                 size=cur_size, bold=cur_bold, italic=cur_italic)
        elif m.group(6): # $math$
            math_content = g0[1:-1]   # strip $
            # Plain numeric values (e.g. $-0,92$) are better as normal text
            # so they match the surrounding table/text formatting.
            if re.fullmatch(r"\s*[-+]?\d+(?:[.,]\d+)?\s*", math_content):
                plain_num = math_content.strip().replace(".", ",")
                _run(para, plain_num, size=cur_size, bold=cur_bold, italic=cur_italic)
            else:
                omml = None
                if TEMP_DIR and PANDOC_PATH:
                    try:
                        omml = latex_to_omml_element(g0, TEMP_DIR, PANDOC_PATH)
                    except Exception:
                        pass
                if omml is not None:
                    para._p.append(omml)
                else:
                    _run(para, _clean(strip_fmt(math_content)),
                         italic=True, size=cur_size, bold=cur_bold)
        elif m.group(7): # \label, \ref, \pageref — skip
            pass
        elif m.group(8): # font switches: \bfseries, \normalfont, \itshape, \large, etc.
            cmd = g0.lstrip("\\").strip()
            if cmd == "bfseries":
                cur_bold = True
            elif cmd == "normalfont":
                cur_bold = False
                cur_italic = False
            elif cmd == "itshape":
                cur_italic = True
            elif cmd in size_map:
                cur_size = size_map[cmd]
            # Other unrecognized switches become no-ops
        elif m.group(9): # generic \cmd
            plain = strip_fmt(g0)
            if plain:
                _run(para, _clean(plain), size=cur_size, bold=cur_bold, italic=cur_italic)
        else:            # plain text
            cleaned = _clean(g0)
            if cleaned:
                _run(para, cleaned, size=cur_size, bold=cur_bold, italic=cur_italic)


# ─────────────────────────────────────────────────────────────────────────────
# Preamble parser
# ─────────────────────────────────────────────────────────────────────────────

def _extract_simple_braced(text: str, cmd: str) -> str | None:
    r"""Extract content of \cmd{...} allowing one level of nested braces."""
    idx = text.find(f"\\{cmd}{{")
    if idx == -1:
        return None
    start = idx + len(f"\\{cmd}{{") - 1
    content = _extract_balanced_braces(text, start)
    return content


def _detect_documentclass(preamble: str) -> tuple[str, list[str]]:
    """Return (class_name, options_list) from \\documentclass[...]{...}."""
    m = re.search(r"\\documentclass\[(.*?)\]\{(\w+)\}", preamble, re.DOTALL)
    if not m:
        return "", []
    opts = [o.strip() for o in m.group(1).split(",")]
    return m.group(2), opts


def _apply_documentclass_to_config(docclass: str, opts: list[str], config: Config):
    """Infer page size and base font size from document class options."""
    # Page size
    if "letterpaper" in opts:
        config.page_width_inches = 8.5
        config.page_height_inches = 11.0
    elif "legalpaper" in opts:
        config.page_width_inches = 8.5
        config.page_height_inches = 14.0
    elif "a4paper" in opts or "a4" in opts:
        config.page_width_inches = 8.27
        config.page_height_inches = 11.69
    elif "a5paper" in opts or "a5" in opts:
        config.page_width_inches = 5.83
        config.page_height_inches = 8.27

    # Font size
    for opt in opts:
        if opt in ("10pt", "11pt", "12pt"):
            config.font_size_normal = int(opt.replace("pt", ""))
            break

    # Default margins depend on class side
    if docclass in ("report", "book"):
        # reasonable defaults if geometry is not present
        config.margin_left = Cm(2.54)
        config.margin_right = Cm(2.54)
        config.margin_top = Cm(2.5)
        config.margin_bottom = Cm(2.0)
    else:
        config.margin_left = Cm(2.54)
        config.margin_right = Cm(2.54)
        config.margin_top = Cm(2.5)
        config.margin_bottom = Cm(2.0)


def _apply_language_to_config(preamble: str, config: Config):
    """Detect babel/polyglossia language and set captions accordingly."""
    lang = None
    m = re.search(r"\\usepackage\[(.*?)\]\{babel\}", preamble, re.DOTALL)
    if m:
        opts = [o.strip().lower() for o in m.group(1).split(",")]
        for o in opts:
            if o in ("spanish", "english", "french", "german", "portuguese"):
                lang = o
                break
    if not lang:
        m = re.search(r"\\setmainlanguage\{(.*?)\}", preamble, re.DOTALL)
        if m:
            lang = m.group(1).strip().lower()

    if lang == "spanish":
        config.language = "es"
        config.table_name = "Tabla"
        config.figure_name = "Figura"
        config.chapter_prefix = "Capítulo "
        config.appendix_prefix = "Apéndice "
        config.bib_title_default = "Referencias"
        config.toc_title = "ÍNDICE DE CONTENIDOS"
        config.lof_title = "ÍNDICE DE FIGURAS"
        config.lot_title = "ÍNDICE DE TABLAS"
        config.toc_placeholder = "Haga clic derecho y seleccione Actualizar campo para generar."
        config.lof_placeholder = "No hay entradas de figuras. Haga clic derecho y seleccione Actualizar campo."
        config.lot_placeholder = "No hay entradas de tablas. Haga clic derecho y seleccione Actualizar campo."
    elif lang == "english":
        config.language = "en"
        config.table_name = "Table"
        config.figure_name = "Figure"
        config.chapter_prefix = "Chapter "
        config.appendix_prefix = "Appendix "
        config.bib_title_default = "References"
        config.toc_title = "TABLE OF CONTENTS"
        config.lof_title = "LIST OF FIGURES"
        config.lot_title = "LIST OF TABLES"
        config.toc_placeholder = "Right-click and select Update Field to generate."
        config.lof_placeholder = "No figure entries found. Right-click and select Update Field to generate."
        config.lot_placeholder = "No table entries found. Right-click and select Update Field to generate."

    # Override with explicit \figurename / \tablename
    fig_name = _extract_simple_braced(preamble, "figurename")
    if fig_name:
        config.figure_name = strip_fmt(fig_name)
    tab_name = _extract_simple_braced(preamble, "tablename")
    if tab_name:
        config.table_name = strip_fmt(tab_name)


def parse_preamble(src: str, config: Config | None = None) -> dict:
    r"""
    Extract \newcommand definitions and \setcounter values.
    Populates global MACROS dict and the provided Config object.
    Returns counters dict.
    """
    if config is None:
        config = CONFIG

    counters = {}

    # \setcounter{name}{value}
    for m in re.finditer(r"\\setcounter\{(\w+)\}\{(\d+)\}", src):
        counters[m.group(1)] = int(m.group(2))

    plazo = counters.get("PlazoTotal", 60)

    # Document class and language drive defaults
    docclass, opts = _detect_documentclass(src)
    _apply_documentclass_to_config(docclass, opts, config)
    _apply_language_to_config(src, config)

    # \newcommand{\Name}{Definition}  — single-level braces
    for m in re.finditer(
        r"\\newcommand\{(\\[A-Za-z]+)\}\{((?:[^{}]|\{[^{}]*\})*)\}",
        src
    ):
        cmd, dfn = m.group(1), m.group(2)
        MACROS[cmd + "{}"] = dfn
        MACROS[cmd]        = dfn

    # \newglossaryentry{key}{name={name},description={...}}  -> macros \gls{key} and \Gls{key}
    pos = 0
    gls_keys = set()
    while True:
        idx = src.find("\\newglossaryentry{", pos)
        if idx == -1:
            break
        key_start = idx + len("\\newglossaryentry{") - 1  # points to '{' before key
        key = _extract_balanced_braces(src, key_start)
        if key is None:
            pos = key_start
            continue
        block_start = key_start + len(key) + 2
        if block_start >= len(src) or src[block_start] != "{":
            pos = block_start
            continue
        block = _extract_balanced_braces(src, block_start)
        if block is None:
            pos = block_start
            continue
        name_idx = block.find("name=")
        if name_idx != -1:
            name_start = name_idx + len("name=")
            if name_start < len(block) and block[name_start] == "{":
                name = _extract_balanced_braces(block, name_start)
                if name is not None:
                    MACROS[f"\\gls{{{key}}}"] = name
                    gls_keys.add(f"\\gls{{{key}}}")
                    if name:
                        cap_name = name[0].upper() + name[1:]
                        MACROS[f"\\Gls{{{key}}}"] = cap_name
                        gls_keys.add(f"\\Gls{{{key}}}")
        pos = block_start + len(block) + 2

    # Month / Year / generic placeholders
    date_info = {
        "\\MONTH":  "",
        "\\YEAR":   "",
        "\\month":  "",
        "\\year":   "",
    }
    for k, v in date_info.items():
        MACROS[k]       = v
        MACROS[k + "{}"] = v

    # Milestone day commands (arithmetic on PlazoTotal)
    milestones = {
        "\\DiaI":     str((8  * plazo + 50) // 100),
        "\\DiaII":    str((25 * plazo + 50) // 100),
        "\\DiaIII":   str((50 * plazo + 50) // 100),
        "\\DiaIV":    str((75 * plazo + 50) // 100),
        "\\DiaFinal": str(plazo),
        "\\arabic{PlazoTotal}": str(plazo),
    }
    for k, v in milestones.items():
        MACROS[k]       = v
        MACROS[k + "{}"] = v

    # Recursively resolve nested macro refs (one pass)
    for key in list(MACROS.keys()):
        for k2, v2 in list(MACROS.items()):
            MACROS[key] = MACROS[key].replace(k2, v2)

    # Strip remaining LaTeX from macro values (preserve glossary names so $math$ stays intact)
    for key in list(MACROS.keys()):
        if key not in gls_keys:
            MACROS[key] = strip_fmt(MACROS[key])

    return counters


# ─────────────────────────────────────────────────────────────────────────────
# Document setup
# ─────────────────────────────────────────────────────────────────────────────

def _parse_length(value: str):
    """Convert a LaTeX length like '2.0cm', '1.5in', '10pt' to a docx length object."""
    value = value.strip()
    m = re.match(r"([\d.]+)\s*([a-zA-Z]+)", value)
    if not m:
        return None
    num = float(m.group(1))
    unit = m.group(2).lower()
    if unit == "cm":
        return Cm(num)
    elif unit == "mm":
        return Mm(num)
    elif unit == "in":
        return Inches(num)
    elif unit == "pt":
        return Pt(num)
    elif unit == "em":
        # Approximate em as 12 pt (matches LaTeX default)
        return Pt(num * 12)
    elif unit == "ex":
        return Pt(num * 5)
    return None


def _find_executable(name: str, windows_hints: list[str] | None = None) -> str:
    """Find an executable: first in PATH, then in common Windows locations."""
    # 1. PATH
    found = shutil.which(name)
    if found:
        return found

    # 2. Common Windows locations
    if sys.platform == "win32" and windows_hints:
        for p in windows_hints:
            if Path(p).exists():
                return p

    return name


def parse_geometry(preamble: str, config: Config | None = None) -> dict:
    """Extract margin values from \\usepackage[...]{geometry} and update config.

    Returns the raw margins dict for backwards compatibility; the canonical
    values are written directly into config.
    """
    if config is None:
        config = CONFIG

    geom_m = re.search(r"\\usepackage\[(.*?)\]\{geometry\}", preamble, re.DOTALL)
    margins = {}
    if geom_m:
        opts = geom_m.group(1)
        for opt in re.split(r",\s*", opts):
            if "=" not in opt:
                continue
            key, val = opt.split("=", 1)
            key = key.strip().lower()
            val = val.strip()
            if key in ("left", "lmargin"):
                margins["left"] = _parse_length(val)
            elif key in ("right", "rmargin"):
                margins["right"] = _parse_length(val)
            elif key in ("top", "tmargin"):
                margins["top"] = _parse_length(val)
            elif key in ("bottom", "bmargin"):
                margins["bottom"] = _parse_length(val)
            elif key == "margin":
                length = _parse_length(val)
                if length is not None:
                    margins.setdefault("left", length)
                    margins.setdefault("right", length)
                    margins.setdefault("top", length)
                    margins.setdefault("bottom", length)

    # Apply to config, falling back to existing config defaults
    config.margin_left   = margins.get("left",   config.margin_left)
    config.margin_right  = margins.get("right",  config.margin_right)
    config.margin_top    = margins.get("top",    config.margin_top)
    config.margin_bottom = margins.get("bottom", config.margin_bottom)

    # Calculate text body dimensions from page size and margins
    def _to_inches(length):
        if length is None:
            return 0.0
        if isinstance(length, Inches):
            return length.inches
        if isinstance(length, Cm):
            return length.cm / 2.54
        if isinstance(length, Mm):
            return length.mm / 25.4
        if isinstance(length, Pt):
            return length.pt / 72.0
        return 0.0

    config.text_width_inches = (
        config.page_width_inches
        - _to_inches(config.margin_left)
        - _to_inches(config.margin_right)
    )
    config.text_height_inches = (
        config.page_height_inches
        - _to_inches(config.margin_top)
        - _to_inches(config.margin_bottom)
    )
    return margins


def setup_document(doc: Document, config: Config | None = None):
    """Apply page size, margins, and base styles from config."""
    if config is None:
        config = CONFIG

    sec = doc.sections[0]
    sec.page_width    = Inches(config.page_width_inches)
    sec.page_height   = Inches(config.page_height_inches)
    sec.top_margin    = config.margin_top
    sec.bottom_margin = config.margin_bottom
    sec.left_margin   = config.margin_left
    sec.right_margin  = config.margin_right

    doc.styles["Normal"].font.name = config.font_name
    doc.styles["Normal"].font.size = Pt(config.font_size_normal)
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)

    base = config.font_size_normal
    heading_cfg = {
        "Heading 1": (base + 3, True,  True),
        "Heading 2": (base + 2, True,  False),
        "Heading 3": (base + 1, True,  False),
        "Heading 4": (base,     False, True),
        "Heading 5": (base,     False, True),
    }
    for name, (sz, bold, is_caps) in heading_cfg.items():
        try:
            s = doc.styles[name]
            s.font.name  = config.font_name
            s.font.size  = Pt(sz)
            s.font.bold  = bold
            s.font.color.rgb = C_BLACK
            if is_caps:
                s.font.all_caps = True
        except Exception:
            pass


def add_header(doc: Document, logo_path: Path, right_text: str):
    """Left: logo image. Right: project name. With horizontal line below."""
    sec    = doc.sections[0]
    header = sec.header
    header.is_linked_to_previous = False
    
    # Clear any existing paragraphs
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    
    # Create table with 2 columns spanning full width
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    
    # Left cell: Logo
    left_cell = table.cell(0, 0)
    left_cell.width = Inches(1.5)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_para.paragraph_format.space_after = Pt(0.2)  # Small space after
    
    if logo_path and logo_path.exists():
        run_img = left_para.add_run()
        run_img.add_picture(str(logo_path), height=Cm(0.8))
    
    # Right cell: Text aligned to right
    right_cell = table.cell(0, 1)
    right_cell.width = Inches(5.0)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_after = Pt(0.2)  # Small space after
    _run(right_para, right_text, size=PT_SMALL)
    
    # Add bottom border to last row's cells
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        # Only bottom border
        for edge in ['top', 'left', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), '808080')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)


def add_footer(doc: Document):
    """Centred page number."""
    sec    = doc.sections[0]
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    _add_page_number_field(run)


# ─────────────────────────────────────────────────────────────────────────────
# Title page
# ─────────────────────────────────────────────────────────────────────────────

def add_title_page(doc: Document, title_img_path: Path | None,
                   title_img_width: float | None, title_img_height: float | None,
                   title: str, author: str, date_str: str):
    # Optional title image (e.g. \includegraphics inside \title)
    if title_img_path and title_img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kwargs = {}
        if title_img_width is not None:
            kwargs["width"] = Inches(title_img_width)
        if title_img_height is not None:
            kwargs["height"] = Inches(title_img_height)
        if not kwargs:
            kwargs["width"] = Inches(4)
        p.add_run().add_picture(str(title_img_path), **kwargs)
        doc.add_paragraph()

    # Main title text
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, title, bold=True, size=16)
        doc.add_paragraph()

    # Author
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, author, bold=True, size=PT_NORM)
        doc.add_paragraph()

    # Date
    if date_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, date_str, bold=True, size=PT_NORM)


# ─────────────────────────────────────────────────────────────────────────────
# Environment extractor
# ─────────────────────────────────────────────────────────────────────────────

def _extract_balanced_braces(text: str, start: int):
    """Extract content inside balanced braces starting at `start` (which must be '{')."""
    if start >= len(text) or text[start] != "{":
        return None
    depth = 1
    i = start + 1
    while i < len(text) and depth > 0:
        if text[i] == "{":
            depth += 1
        elif text[i] == "}":
            depth -= 1
        i += 1
    if depth == 0:
        return text[start + 1 : i - 1]
    return None


def _extract_caption_content(text: str) -> str | None:
    """Extract the last \\caption{...} content, handling nested braces."""
    idx = text.rfind("\\caption{")
    if idx == -1:
        return None
    brace_start = idx + len("\\caption{") - 1  # position of '{'
    content = _extract_balanced_braces(text, brace_start)
    return content


def _remove_captions(text: str) -> str:
    """Remove all \\caption{...} occurrences, handling nested braces."""
    result = []
    pos = 0
    while True:
        idx = text.find("\\caption{", pos)
        if idx == -1:
            result.append(text[pos:])
            break
        result.append(text[pos:idx])
        brace_start = idx + len("\\caption{") - 1
        content = _extract_balanced_braces(text, brace_start)
        if content is None:
            result.append(text[idx:idx + len("\\caption")])
            pos = idx + len("\\caption")
            continue
        after = brace_start + len(content) + 2
        # skip trailing \\ if present
        if text[after:after + 2] == "\\\\":
            after += 2
        pos = after
    return "".join(result)


def extract_env(text: str, env: str, start: int = 0):
    r"""
    Find first balanced \begin{env}...\end{env} starting at `start`.
    Returns (abs_begin, abs_end, inner) or None.
    """
    bp = re.compile(r"\\begin\{" + re.escape(env) + r"\*?\}")
    ep = re.compile(r"\\end\{"   + re.escape(env) + r"\*?\}")

    m0 = bp.search(text, start)
    if not m0:
        return None

    depth = 1
    pos   = m0.end()
    while depth > 0 and pos < len(text):
        nb = bp.search(text, pos)
        ne = ep.search(text, pos)
        if ne is None:
            break
        if nb and nb.start() < ne.start():
            depth += 1
            pos = nb.end()
        else:
            depth -= 1
            if depth == 0:
                return (m0.start(), ne.end(), text[m0.end(): ne.start()])
            pos = ne.end()
    return None


def _strip_column_spec(text: str):
    """
    Remove the first balanced {col_spec} block from text.
    Returns (col_spec_content, rest_of_text).
    Handles nested braces like p{0.06\\textwidth}.
    """
    text = text.lstrip()
    if not text.startswith("{"):
        return "", text
    depth = 0
    for i, c in enumerate(text):
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return text[1:i], text[i + 1:]
    return "", text


def _flatten_single_cell_tabulars(content: str) -> str:
    """Replace nested single-cell tabulars with their inner text.

    Authors often wrap a cell header in \\begin{tabular}[c]{@{}c@{}}text
    \\end{tabular} to centre it. Pandoc leaves the wrapper visible, so we
    flatten it when the inner content is clearly a single cell (no row or
    column separators). The regex only matches innermost tabulars so it
    never accidentally consumes the outer table.
    """
    # Match an innermost tabular: its body cannot contain another
    # \begin{tabular} or \end{tabular}. The column spec may itself contain
    # nested braces (e.g. {@{}c@{}}), so we allow one level of nesting.
    pattern = re.compile(
        r"\\begin\{tabular\}(?:\[[^\]]*\])?"
        r"\{(?:[^{}]|\{[^{}]*\})*\}\s*"
        r"((?:(?!\\begin\{tabular\}|\\end\{tabular\}).)*?)"
        r"\\end\{tabular\}",
        re.DOTALL
    )

    changed = True
    while changed:
        changed = False

        def replace(match: re.Match) -> str:
            nonlocal changed
            inner = match.group(1).strip()
            # If it looks like a real nested table, keep it.
            if "\\\\" in inner or "&" in inner:
                return match.group(0)
            changed = True
            return inner

        content = pattern.sub(replace, content)

    return content


def _parse_col_widths_dxa(col_spec: str, n_cols: int,
                           content_dxa: int = 9026) -> list:
    """
    Parse a LaTeX column spec and return column widths in DXA units.
    Handles p{frac\\textwidth}, X, l, r, c.
    Unknown/flexible columns share the remaining width equally.
    """
    fixed_fracs = []   # (index, fraction)
    flex_indices = []

    # Extract individual column descriptors (ignore | and @{...})
    col_spec_clean = re.sub(r"@\{[^}]*\}", "", col_spec)
    col_spec_clean = col_spec_clean.replace("|", "")

    tokens = re.findall(
        r"p\{([\d.]+)\\textwidth\}"   # p{0.XX\textwidth}
        r"|m\{([\d.]+)\\textwidth\}"  # m{0.XX\textwidth}
        r"|b\{([\d.]+)\\textwidth\}"  # b{0.XX\textwidth}
        r"|X"                          # X (tabularx flexible)
        r"|[lrcL]",                    # l r c or L
        col_spec_clean
    )

    cols = []
    for tok in tokens:
        if isinstance(tok, tuple):
            frac = tok[0] or tok[1] or tok[2]
            if frac:
                cols.append(float(frac))
            else:
                cols.append(None)  # X / l / r / c
        else:
            cols.append(None)

    # If we parsed wrong number of cols, fall back to equal widths
    if not cols or len(cols) != n_cols:
        equal = content_dxa // n_cols if n_cols else content_dxa
        return [equal] * n_cols

    fixed_total = sum(f for f in cols if f is not None)
    flex_count  = sum(1 for f in cols if f is None)
    remaining   = content_dxa - int(fixed_total * content_dxa)
    flex_each   = (remaining // flex_count) if flex_count else 0

    result = []
    for f in cols:
        if f is not None:
            result.append(int(f * content_dxa))
        else:
            result.append(max(flex_each, 500))  # minimum 500 DXA

    # Adjust rounding to match content width exactly
    diff = content_dxa - sum(result)
    if result:
        result[-1] += diff
    return result



def _is_rule(line: str) -> bool:
    return bool(re.match(r"\\(toprule|midrule|bottomrule|hline)\b", line.strip()))


def _is_category_row(line: str) -> bool:
    r"""Detect rows that are category/section headers (\multicolumn spanning all cols)."""
    return bool(re.match(r"\s*\\multicolumn\b", line.strip()) and "&" not in line)


def _expand_multicolumn(line: str) -> str:
    r"""
    Replace \multicolumn{n}{align}{text} with (text)(& repeat n-1 times).
    """
    def repl(m):
        n    = int(m.group(1))
        text = m.group(3)
        return text + " & " * (n - 1)

    return re.sub(
        r"\\multicolumn\{(\d+)\}\{([^}]*)\}\{((?:[^{}]|\{[^{}]*\})*)\}",
        repl, line
    )


# Contador global de tablas para usar tablas Pandoc pre-generadas
TABLE_COUNTER = [0]

# Contador global de figuras tikz
TIKZ_COUNTER = [0]

# Contador global de figuras (para captions numerados)
FIGURE_COUNTER = [0]

# Preámbulo LaTeX original (se establece en parse_body)
PREAMBLE_GLOBAL: str = ""

# Rutas de búsqueda de imágenes extraídas de \graphicspath{{...}}
GRAPHICSPATHS: list[Path] = []

# Opciones por defecto de enumitem extraídas del preámbulo (se establece en parse_body)
# Formato: {nivel: label_format, None: label_format_global}
ENUMITEM_DEFAULTS: dict[int | None, str] = {}

def apply_booktabs_style(table):
    """
    Aplica estilo LaTeX booktabs a una tabla:
    - Sin bordes verticales
    - Líneas horizontales: top (gruesa), mid (media entre filas), bottom (gruesa)
    """
    rows = list(table.rows)
    if not rows:
        return
    
    num_rows = len(rows)
    
    for row_idx, row in enumerate(rows):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            # Sin bordes verticales (left/right)
            for edge in ['left', 'right']:
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            
            # Top border: primera fila = toprule (grueso)
            if row_idx == 0:
                b = OxmlElement('w:top')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '12')  # Grueso
                b.set(qn('w:color'), '000000')
                tcBorders.append(b)
            else:
                # Otras filas = sin línea arriba (la línea viene de la celda de arriba)
                b = OxmlElement('w:top')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            
            # Bottom border:
            if row_idx == num_rows - 1:
                # Última fila = bottomrule (grueso)
                b = OxmlElement('w:bottom')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '12')
                b.set(qn('w:color'), '000000')
                tcBorders.append(b)
            elif row_idx == 0:
                # Después del header = midrule (media)
                b = OxmlElement('w:bottom')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '6')
                b.set(qn('w:color'), '000000')
                tcBorders.append(b)
            else:
                # Entre filas = línea ligera (cmidrule style)
                b = OxmlElement('w:bottom')
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '4')  # Más delgada
                b.set(qn('w:color'), 'C0C0C0')  # Gris claro
                tcBorders.append(b)
            
            tcPr.append(tcBorders)
            
            # Asegurar que el texto tenga fuente Calibri y color negro
            for para in cell.paragraphs:
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = FONT
                    run.font.color.rgb = C_BLACK
                    if run.font.size is None:
                        run.font.size = Pt(PT_SMALL)

def normalize_text(text: str) -> str:
    """Normaliza texto para comparación (quita espacios extras, lowercase)."""
    return ' '.join(text.strip().lower().split())




# ─────────────────────────────────────────────────────────────────────────────
# enumitem label helpers
# ─────────────────────────────────────────────────────────────────────────────

def _int_to_roman(num: int, upper: bool = True) -> str:
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syms = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
    res = ''
    for v, s in zip(val, syms):
        while num >= v:
            res += s
            num -= v
    return res if upper else res.lower()


def _int_to_alpha(num: int, upper: bool = True) -> str:
    res = ''
    while num > 0:
        num -= 1
        res = chr(ord('A') + (num % 26)) + res
        num //= 26
    return res if upper else res.lower()


def _format_enumitem_label(label_format: str, index: int) -> str:
    if not label_format:
        return str(index)
    fmt = label_format
    fmt = re.sub(r'\\roman\*', lambda m: _int_to_roman(index, upper=False), fmt)
    fmt = re.sub(r'\\Roman\*', lambda m: _int_to_roman(index, upper=True), fmt)
    fmt = re.sub(r'\\alph\*',  lambda m: _int_to_alpha(index, upper=False), fmt)
    fmt = re.sub(r'\\Alph\*',  lambda m: _int_to_alpha(index, upper=True), fmt)
    fmt = re.sub(r'\\arabic\*', lambda m: str(index), fmt)
    return fmt


def _extract_optional_arg(text: str):
    text = text.lstrip()
    if not text.startswith('['):
        return None, text
    depth = 0
    i = 0
    while i < len(text):
        c = text[i]
        if c == '[':
            depth += 1
        elif c == ']':
            depth -= 1
            if depth == 0:
                arg = text[1:i]
                rest = text[i + 1:].lstrip()
                return arg, rest
        elif c == "\\" and i + 1 < len(text):
            i += 1
        i += 1
    return None, text


def _parse_enumitem_label(optional_arg: str) -> str:
    if not optional_arg:
        return None
    m = re.search(r'(?:^|[,\[]\s*)label\s*=\s*(.+?)(?:\s*,\s*\w+\s*=|$)', optional_arg, re.DOTALL)
    if not m:
        return None
    label = m.group(1).strip()
    # Strip trailing value-less enumitem options (e.g. nosep, noitemsep, wide)
    # that the regex may have captured because they have no '='.
    label = re.sub(
        r'\s*,\s*(nosep|noitemsep|wide|leftmargin|labelsep|itemsep|topsep|parsep|partopsep|align|font|format|before|after|ref)\s*$',
        '', label, flags=re.IGNORECASE
    )
    return label


def _extract_enumitem_defaults(preamble: str) -> dict[int | None, str]:
    r"""Extrae los labels por defecto de \setlist[enumerate]{...} del pre�mbulo."""
    defaults: dict[int | None, str] = {}
    # Busca \setlist[enumerate,...]{...} o \setlist[enumerate]{...}
    for m in re.finditer(r'\\setlist\s*\[\s*enumerate\s*(?:,\s*(\d+))?\s*\]\s*\{', preamble):
        level = int(m.group(1)) if m.group(1) else None
        content = _extract_balanced_braces(preamble, m.end() - 1)
        if content is None:
            continue
        label = _parse_enumitem_label(content)
        if label:
            defaults[level] = label
    return defaults


def _get_enumitem_default(level: int) -> str | None:
    """Devuelve el label por defecto para el nivel dado, o el global si no hay específico."""
    if level in ENUMITEM_DEFAULTS:
        return ENUMITEM_DEFAULTS[level]
    return ENUMITEM_DEFAULTS.get(None)


def is_duplicate_row(row_texts: list, seen_rows: list) -> bool:
    """
    Detecta si una fila es idéntica a una ya vista.
    Compara el contenido normalizado de todas las celdas.
    """
    normalized = tuple(normalize_text(t) for t in row_texts)
    
    for seen in seen_rows:
        if normalized == seen:
            return True
    return False


def is_continuation_row_cells(cells_texts: list) -> bool:
    """
    Detecta si una fila es de 'continued' o 'continues on next page'.
    Estas filas vienen de longtable y no queremos mostrarlas.
    Recibe una lista de textos de celdas.
    """
    for text in cells_texts:
        text = text.strip().lower()
        # Patrones de filas de continuación
        if any(pattern in text for pattern in [
            'continued',
            'continues on next page',
            'table continued',
            'continúa en la siguiente página',
            'continua en la siguiente página',
            '(continued)'
        ]):
            return True
    return False


def is_continuation_row(row) -> bool:
    """Versión para filas de tabla de docx."""
    texts = [cell.text for cell in row.cells]
    return is_continuation_row_cells(texts)


def compile_tikz_to_png(tikz_code: str, preamble: str, temp_dir: Path, fig_num: int):
    """
    Compila un bloque tikzpicture a PNG mediante pdflatex + pdftoppm.
    Recorta el espacio en blanco sobrante con Pillow.
    Devuelve la ruta del PNG o None si falla.
    """
    # Limpiar preámbulo: quitar \documentclass y \usepackage{geometry} conflictivo
    clean_preamble = re.sub(r"\\documentclass\[.*?\]\{.*?\}\s*", "", preamble, flags=re.DOTALL)
    clean_preamble = re.sub(r"\\usepackage\[margin=[^\]]*\]\{geometry\}\s*", "", clean_preamble)

    tex_content = (
        r"\documentclass{article}" + "\n"
        r"\usepackage[margin=0pt]{geometry}" + "\n"
        r"\pagestyle{empty}" + "\n"
        + clean_preamble + "\n"
        r"\begin{document}" + "\n"
        r"\noindent" + "\n"
        + tikz_code + "\n"
        r"\end{document}" + "\n"
    )

    tex_file = temp_dir / f"tikz_{fig_num:02d}.tex"
    tex_file.write_text(tex_content, encoding="utf-8")

    pdflatex = CONFIG.pdflatex_path
    try:
        result = subprocess.run(
            [pdflatex, "-interaction=nonstopmode", str(tex_file.name)],
            cwd=str(temp_dir),
            capture_output=True,
            text=True,
            timeout=CONFIG.timeout_pdflatex,
        )
    except Exception as e:
        print(f"  [WARN] pdflatex falló para tikz {fig_num}: {e}")
        return None

    pdf_file = temp_dir / f"tikz_{fig_num:02d}.pdf"
    if not pdf_file.exists() or pdf_file.stat().st_size == 0:
        print(f"  [WARN] PDF no generado para tikz {fig_num}")
        return None

    pdftoppm = CONFIG.pdftoppm_path
    png_file = temp_dir / f"tikz_{fig_num:02d}.png"
    try:
        subprocess.run(
            [pdftoppm, "-png", "-r", str(CONFIG.image_tikz_dpi), "-singlefile", str(pdf_file), str(temp_dir / f"tikz_{fig_num:02d}")],
            capture_output=True,
            check=False,
            timeout=CONFIG.timeout_pdftoppm,
        )
    except Exception as e:
        print(f"  [WARN] pdftoppm falló para tikz {fig_num}: {e}")
        return None

    if not png_file.exists():
        return None

    # Recortar espacio en blanco con Pillow
    try:
        from PIL import Image
        img = Image.open(str(png_file))
        if img.mode in ("RGBA", "LA", "P"):
            img = img.convert("RGB")
        gray = img.convert("L")
        # Invertir para que el blanco sea 0 y lo demás >0
        bbox = gray.point(lambda x: 255 - x).getbbox()
        if bbox:
            img = img.crop(bbox)
            img.save(str(png_file))
    except Exception as e:
        print(f"  [WARN] No se pudo recortar tikz {fig_num}: {e}")

    return png_file


def _parse_graphics_size(opts_str: str, img_path: Path):
    r"""
    Parse \\includegraphics[...] options and return (width_inches, height_inches).
    Supports: width=0.65\\textwidth, width=\\textwidth, width=5cm, height=0.28\\textheight,
              height=\\textheight, height=3cm, scale=0.5, keepaspectratio.
    Reads natural image size for aspect ratio if needed.
    """
    width = None
    height = None
    scale = None
    keep_aspect = "keepaspectratio" in opts_str.lower().replace("-", "").replace(" ", "")

    # width = fraction \textwidth (fraction optional, defaults to 1.0)
    m = re.search(r"width=([\d.]*)\\?textwidth", opts_str)
    if m:
        frac = float(m.group(1)) if m.group(1) else 1.0
        width = frac * TEXT_WIDTH_INCHES
    else:
        m = re.search(r"width=([\d.]*)\\?linewidth", opts_str)
        if m:
            frac = float(m.group(1)) if m.group(1) else 1.0
            width = frac * TEXT_WIDTH_INCHES
        else:
            m = re.search(r"width=([\d.]+)cm", opts_str)
            if m:
                width = float(m.group(1)) / 2.54
            else:
                m = re.search(r"width=([\d.]+)mm", opts_str)
                if m:
                    width = float(m.group(1)) / 25.4
                else:
                    m = re.search(r"width=([\d.]+)in", opts_str)
                    if m:
                        width = float(m.group(1))

    # height = fraction \textheight (fraction optional, defaults to 1.0)
    m = re.search(r"height=([\d.]*)\\?textheight", opts_str)
    if m:
        frac = float(m.group(1)) if m.group(1) else 1.0
        height = frac * TEXT_HEIGHT_INCHES
    else:
        m = re.search(r"height=([\d.]+)cm", opts_str)
        if m:
            height = float(m.group(1)) / 2.54
        else:
            m = re.search(r"height=([\d.]+)mm", opts_str)
            if m:
                height = float(m.group(1)) / 25.4
            else:
                m = re.search(r"height=([\d.]+)in", opts_str)
                if m:
                    height = float(m.group(1))

    # scale
    m = re.search(r"scale=([\d.]+)", opts_str)
    if m:
        scale = float(m.group(1))

    # Natural size for aspect ratio / scale
    try:
        from PIL import Image
        img = Image.open(str(img_path))
        px_w, px_h = img.size
        nat_w = px_w / 96.0
        nat_h = px_h / 96.0
    except Exception:
        nat_w = 4.0
        nat_h = 3.0

    if scale is not None:
        width = nat_w * scale
        height = nat_h * scale
    elif width is not None and height is None:
        ratio = width / nat_w
        height = nat_h * ratio
    elif height is not None and width is None:
        ratio = height / nat_h
        width = nat_w * ratio
    elif width is not None and height is not None:
        # Both width and height were specified
        if keep_aspect:
            # Fit inside the width x height box while preserving aspect ratio
            scale_w = width / nat_w
            scale_h = height / nat_h
            scale = min(scale_w, scale_h)
            width = nat_w * scale
            height = nat_h * scale
        # Without keepaspectratio LaTeX distorts the image, so we keep
        # the explicit width and height as-is.

    return width, height


def _extract_graphicspaths(source: str) -> list[Path]:
    r"""Extrae rutas de \graphicspath{{dir1/}{dir2/}...} del preámbulo."""
    paths = []
    for m in re.finditer(r"\\graphicspath\s*\{((?:\s*\{[^}]*\}\s*)*)\}", source):
        block = m.group(1)
        for pm in re.finditer(r"\{([^}]*)\}", block):
            p = pm.group(1).strip().replace("\\", "/").strip("/")
            if p:
                paths.append(Path(p))
    return paths


def resolve_image_path(img_name: str, base_dir: Path) -> Path | None:
    """Busca una imagen respetando GRAPHICSPATHS y extensiones comunes."""
    name = img_name.strip().replace("\\", "/")
    # Lista de candidatos: nombre exacto + con extensiones comunes
    candidates = [name]
    if "." not in name.split("/")[-1]:
        for ext in (".png", ".jpg", ".jpeg", ".pdf", ".gif", ".bmp", ".eps"):
            candidates.append(name + ext)
    # Buscar en base_dir y en cada graphicspath relativo a base_dir
    search_dirs = [base_dir] + [base_dir / p for p in GRAPHICSPATHS]
    for cand in candidates:
        for d in search_dirs:
            p = d / cand
            if p.exists():
                return p
    return None


def _find_logo(base_dir: Path, config: Config | None = None) -> Path | None:
    """Find a logo image, optionally using \\logo{...} from the preamble."""
    if config is None:
        config = CONFIG

    # If the preamble defined \logo{...}, try that first
    logo_cmd = _extract_simple_braced(PREAMBLE_GLOBAL, "logo")
    if logo_cmd:
        p = resolve_image_path(logo_cmd, base_dir)
        if p:
            return p

    script_dir = Path(__file__).parent
    for cand in config.logo_candidates:
        for d in (base_dir, script_dir):
            p = d / cand
            if p.exists():
                return p
    return None


def _insert_image_with_size(para, img_path: Path, max_width_inches: float = 6.3, dpi: float = 300.0):
    """Inserta imagen respetando su tamaño natural (basado en píxeles / dpi).
    Si excede el ancho máximo, escala proporcionalmente."""
    try:
        from PIL import Image
        img = Image.open(str(img_path))
        px_w, px_h = img.size
        width = px_w / dpi
        height = px_h / dpi
        if width > max_width_inches:
            ratio = max_width_inches / width
            width = max_width_inches
            height = height * ratio
        para.add_run().add_picture(str(img_path), width=Inches(width), height=Inches(height))
        return True
    except Exception:
        return False


def _render_figure_as_table(doc: Document, inner: str, base_dir: Path) -> bool:
    r"""
    Renderiza una figura con múltiples subfigures en una tabla de Word.

    Soporta dos formatos:
      1) Entornos subfigure modernos:
         \begin{subfigure}[b]{0.32\textwidth}
             \centering
             \includegraphics[width=\textwidth]{...}
             \caption{...}\label{...}
         \end{subfigure}
      2) Comando \subfigure antiguo:
         \subfigure[caption]{\includegraphics[opts]{name}}

    Las filas se detectan por los saltos de línea LaTeX (\\) entre subfigures.
    Las imágenes se escalan para ocupar uniformemente el ancho de texto.
    Devuelve True si se renderizó como tabla, False en caso contrario.
    """

    def _get_natural_size(img_path):
        try:
            from PIL import Image
            img = Image.open(str(img_path))
            px_w, px_h = img.size
            return px_w / 96.0, px_h / 96.0
        except Exception:
            return 4.0, 3.0

    items = []  # dicts: img_name, img_path, opts, caption, start, end

    # ── Formato moderno: \begin{subfigure} ... \end{subfigure} ────────────────
    env_matches = list(re.finditer(r"\\begin\{subfigure\}", inner))
    if env_matches:
        for m in env_matches:
            result = extract_env(inner, "subfigure", m.start())
            if not result:
                continue
            _, end_pos, sf_inner = result
            img_m = re.search(r"\\includegraphics(?:\[([^\]]*)\])?\{([^}]*)\}", sf_inner)
            if not img_m:
                continue
            caption = _extract_caption_content(sf_inner) or ""
            items.append({
                "img_name": img_m.group(2),
                "img_path": resolve_image_path(img_m.group(2), base_dir),
                "opts": img_m.group(1) or "",
                "caption": caption,
                "start": m.start(),
                "end": end_pos,
            })

    # ── Formato antiguo: \subfigure[caption]{\includegraphics...} ────────────
    if not items:
        subfig_pattern = (
            r"\\subfigure\s*\[((?:[^\[\]]|\\\[|\\)*)\]\s*\{+\s*"
            r"\s*\\includegraphics(?:\[([^\]]*)\])?\{([^}]*)\}"
            r"(?:\s*\\label\{[^}]*\})?\s*\}+"
        )
        for sf in re.finditer(subfig_pattern, inner):
            items.append({
                "img_name": sf.group(3),
                "img_path": resolve_image_path(sf.group(3), base_dir),
                "opts": sf.group(2) or "",
                "caption": sf.group(1),
                "start": sf.start(),
                "end": sf.end(),
            })

    if len(items) <= 1:
        return False

    # ── Agrupar en filas usando \\ como separador entre subfigures ───────────
    row_groups = []
    current_row = [items[0]]
    for i in range(1, len(items)):
        segment = inner[items[i - 1]["end"]:items[i]["start"]]
        if re.search(r"\\\\(?:\s*\[[^\]]*\])?", segment):
            row_groups.append(current_row)
            current_row = [items[i]]
        else:
            current_row.append(items[i])
    row_groups.append(current_row)

    max_cols = max(len(g) for g in row_groups)
    if max_cols == 0:
        return False

    # ── Calcular tamaños finales para cada fila ──────────────────────────────
    final_rows = []
    for row in row_groups:
        n_cols = len(row)
        cell_width = TEXT_WIDTH_INCHES / max_cols
        scaled_row = []
        for item in row:
            nat_w, nat_h = _get_natural_size(item["img_path"])
            # Escalar para que el ancho ocupe la celda disponible
            scale = cell_width / nat_w if nat_w > 0 else 1.0
            final_w = cell_width
            final_h = nat_h * scale
            scaled_row.append({
                **item,
                "nat_w": nat_w,
                "nat_h": nat_h,
                "final_w": final_w,
                "final_h": final_h,
            })
        final_rows.append(scaled_row)

    # ── Crear la tabla e insertar imágenes ───────────────────────────────────
    table = doc.add_table(rows=len(row_groups), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    subfig_idx = 0
    for ri, row in enumerate(final_rows):
        n_cols = len(row)
        for ci, item in enumerate(row):
            cell = table.cell(ri, ci)
            cell.text = ""
            cell.width = Inches(item["final_w"])

            p_img = cell.paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if item["img_path"] and item["img_path"].exists():
                try:
                    p_img.add_run().add_picture(
                        str(item["img_path"]),
                        width=Inches(item["final_w"]),
                        height=Inches(item["final_h"]),
                    )
                except Exception:
                    _run(p_img, f"[Figure: {item['img_name']}]", italic=True, size=PT_SMALL)
            else:
                _run(p_img, f"[Figure: {item['img_name']}]", italic=True, size=PT_SMALL)

            caption_text = strip_fmt(item["caption"])
            if caption_text:
                p_cap = cell.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                letter = chr(ord('a') + subfig_idx)
                _run(p_cap, f"({letter}) {caption_text}", italic=True, size=PT_SMALL)

            subfig_idx += 1

        # Celdas vacías sobrantes
        for ci in range(n_cols, max_cols):
            cell = table.cell(ri, ci)
            cell.text = ""

    # ── Quitar bordes ────────────────────────────────────────────────────────
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('top', 'bottom', 'left', 'right'):
                b = OxmlElement(f'w:{edge}')
                b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            tcPr.append(tcBorders)

    return True


def render_table_with_pandoc(doc: Document, tab_inner: str, caption: str = ""):
    """
    Inserta tabla pre-generada por Pandoc y aplica estilo booktabs.
    Las tablas deben estar en TEMP_DIR: tabla_XX_pandoc.docx
    Ajusta automaticamente el tamaño de fuente segun el numero de columnas.
    """
    global TEMP_DIR
    TABLE_COUNTER[0] += 1
    table_num = TABLE_COUNTER[0]
    
    pandoc_table_file = TEMP_DIR / f'tabla_{table_num:02d}_pandoc.docx'
    
    if not pandoc_table_file.exists():
        print(f'  [WARN] No se encuentra tabla Pandoc: {pandoc_table_file}')
        return False
    
    try:
        # Cargar documento con tabla Pandoc
        pandoc_doc = Document(str(pandoc_table_file))
        
        if not pandoc_doc.tables:
            print(f'  [WARN] El archivo no tiene tablas: {pandoc_table_file}')
            return False
        
        # Agregar caption si existe
        if caption:
            add_table_caption(doc, caption, table_num)
        
        # Copiar primera tabla del archivo Pandoc
        source_table = pandoc_doc.tables[0]
        
        # Detectar numero de columnas y ajustar fuente automaticamente
        ncols = len(source_table.columns)
        font_size = PT_SMALL  # Default 9pt
        if ncols > 7:
            font_size = 7  # Muy ancha -> fuente pequena
        elif ncols > 5:
            font_size = 8  # Moderadamente ancha
        
        # Filtrar filas de continuación y duplicados (headers repetidos)
        valid_rows = []
        seen_headers = set()  # Para detectar headers duplicados
        
        for row in source_table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            row_text = ' '.join(row_texts).lower()
            
            # Detectar filas de continuación
            if is_continuation_row_cells(row_texts):
                continue

            # Detectar filas completamente vacías (residuos de \midrule, \endhead, etc.)
            if not any(row_texts):
                continue

            # Detectar headers duplicados (fila completamente idéntica a una vista antes)
            row_hash = tuple(normalize_text(t) for t in row_texts)
            if row_hash in seen_headers:
                continue  # Skip fila duplicada
            seen_headers.add(row_hash)

            valid_rows.append(row)
        
        # Fusionar filas de header consecutivas que sean complementarias.
        # Esto ocurre cuando un header de LaTeX ocupa varias filas y cada celda
        # solo tiene texto en una de esas filas (ej. fila 1: nombres de columnas
        # desde la columna 2; fila 2: nombre de la columna 1).
        if valid_rows:
            merged = [cell.text.strip() for cell in valid_rows[0].cells]
            i = 1
            while i < len(valid_rows):
                next_texts = [cell.text.strip() for cell in valid_rows[i].cells]
                if len(next_texts) != len(merged):
                    break
                # Filas complementarias: en cada columna, al menos una celda está vacía.
                if not all((not a or not b) for a, b in zip(merged, next_texts)):
                    break
                if not any(next_texts):
                    break
                merged = [a or b for a, b in zip(merged, next_texts)]
                valid_rows.pop(i)
            # Aplicar el header fusionado a la primera fila
            for j, cell in enumerate(valid_rows[0].cells):
                if merged[j]:
                    cell.text = merged[j]

        if not valid_rows:
            print(f'  [WARN] Tabla {table_num} solo tiene filas de continuación')
            return False

        # Crear nueva tabla solo con filas válidas
        ncols = len(source_table.columns)
        new_table = doc.add_table(rows=len(valid_rows), cols=ncols)
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Detectar numero de columnas y ajustar fuente automaticamente
        font_size = PT_SMALL  # Default 9pt
        if ncols > 7:
            font_size = 7  # Muy ancha -> fuente pequena
        elif ncols > 5:
            font_size = 8  # Moderadamente ancha
        
        # Copiar contenido celda por celda PRESERVANDO formato
        for i, source_row in enumerate(valid_rows):
            new_row = new_table.rows[i]
            
            # Detectar si esta fila tiene celdas combinadas (mismo texto en todas)
            row_texts = [cell.text.strip() for cell in source_row.cells]
            all_same = len(set(row_texts)) == 1 and row_texts[0] and ncols > 1
            
            if all_same:
                # Merge todas las celdas de esta fila
                merged_cell = new_row.cells[0]
                merged_cell.text = row_texts[0]
                for run in merged_cell.paragraphs[0].runs:
                    run.font.name = FONT
                    run.font.size = Pt(font_size)
                    run.bold = True
                # Merge con la ultima celda
                merged_cell.merge(new_row.cells[-1])
            else:
                # Copiar normalmente
                for j, source_cell in enumerate(source_row.cells):
                    if j < len(new_row.cells):
                        new_cell = new_row.cells[j]
                        new_cell.text = ""
                        
                        for source_para in source_cell.paragraphs:
                            has_math = any(child.tag.startswith(f"{{{MATH_NS}}}") for child in source_para._p)
                            if not source_para.text.strip() and not has_math:
                                continue
                            new_para = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
                            new_para.text = source_para.text
                            new_para.alignment = source_para.alignment
                            # Copy OMML math elements (Pandoc produces these inline)
                            if has_math:
                                for child in source_para._p:
                                    if child.tag.startswith(f"{{{MATH_NS}}}"):
                                        cloned = etree.fromstring(etree.tostring(child))
                                        new_para._p.append(cloned)
                            for source_run in source_para.runs:
                                for run in new_para.runs:
                                    run.font.name = FONT
                                    run.font.size = Pt(font_size)
                                    if source_run.bold:
                                        run.bold = True
                                    if source_run.italic:
                                        run.italic = True
        
        # Aplicar estilo booktabs (solo líneas horizontales)
        apply_booktabs_style(new_table)
        
        print(f'  [OK] Tabla {table_num} insertada ({len(valid_rows)} de {len(source_table.rows)} filas)')
        return True
        
    except Exception as e:
        print(f'  [ERROR] Insertando tabla {table_num}: {e}')
        import traceback
        traceback.print_exc()
        return False


def _is_p_only_table(col_spec: str) -> bool:
    """Return True if the table has a single p/m/b column (no l/c/r/X).

    In single-column p{} tables, \\ inside the cell is a line break rather
    than a row separator, so Pandoc's row splitting produces too many rows.
    """
    if not col_spec:
        return False
    spec_clean = re.sub(r">\s*\{[^}]*\}", "", col_spec)
    spec_clean = re.sub(r"@[^{]*", "", spec_clean)
    spec_clean = re.sub(r"\|+", "", spec_clean)
    p_cols = re.findall(r"[pmb]\{[^}]*\}", spec_clean)
    other_cols = re.findall(r"[lcrX]", spec_clean)
    return len(p_cols) == 1 and not other_cols


def _has_complex_col_spec(col_spec: str) -> bool:
    r"""Detect column spec constructs that Pandoc handles poorly, such as
    >{\centering\arraybackslash}p{...}.
    """
    if not col_spec:
        return False
    return bool(re.search(r">\s*\{[^}]*\}", col_spec))


def render_table(doc: Document, tab_inner: str, caption: str = "",
                 ncols_hint: int = 0, col_widths_dxa: list = None,
                 col_spec: str = "", is_array: bool = False):
    """
    Convert LaTeX tabular inner content to a Word table.
    First tries Pandoc, falls back to manual rendering.
    Tables with inline math ($...$) skip Pandoc because Pandoc places
    OMML on separate paragraphs; fallback manual rendering keeps math inline.
    Vertical p-only tables are also rendered manually because Pandoc treats
    in-cell \\ as row separators.
    """
    p_only = _is_p_only_table(col_spec)
    has_inline_math = bool(re.search(r"\$[^$]+\$", tab_inner))
    complex_spec = _has_complex_col_spec(col_spec)

    # Tablas \begin{array} dentro de math display no tienen archivo Pandoc
    # temporal; se renderizan siempre con fallback manual.
    if is_array:
        TABLE_COUNTER[0] += 1
        if caption:
            add_table_caption(doc, caption, TABLE_COUNTER[0])
    elif not p_only and not has_inline_math and not complex_spec and render_table_with_pandoc(doc, tab_inner, caption):
        return
    elif p_only or has_inline_math or complex_spec:
        # Consume the Pandoc counter so subsequent tables stay in sync
        TABLE_COUNTER[0] += 1

    # Fallback: manual rendering
    if caption and not is_array:
        add_table_caption(doc, caption, TABLE_COUNTER[0])

    # Clean tab_inner
    cleaned_inner = re.sub(r"p\{[^}]*\}", "", tab_inner)
    cleaned_inner = re.sub(r"m\{[^}]*\}", "", cleaned_inner)
    cleaned_inner = re.sub(r"b\{[^}]*\}", "", cleaned_inner)
    # Flatten nested single-cell tabulars used for centred headers.
    cleaned_inner = _flatten_single_cell_tabulars(cleaned_inner)

    # Detect "vertical" tables where every column is p/m/b (no l/c/r/X).
    # In such tables \\ inside a cell is a line break, not a row separator.
    p_only = _is_p_only_table(col_spec)

    # Proteger \shortstack{...}: sus \\ internos son saltos de línea de celda,
    # no separadores de fila. Usamos extracción balanceada para soportar
    # \textbf{...} u otros comandos con llaves dentro del \shortstack.
    shortstack_segments = []
    pos = 0
    while True:
        idx = cleaned_inner.find("\\shortstack{", pos)
        if idx == -1:
            break
        brace_start = idx + len("\\shortstack{") - 1  # posición de '{'
        body = _extract_balanced_braces(cleaned_inner, brace_start)
        if body is None:
            pos = idx + len("\\shortstack{")
            continue
        # Los \\ internos se protegen para que no partan la fila
        protected_body = body.replace("\\\\", "\x00SJ\x00")
        shortstack_segments.append(protected_body)
        replacement = f"\\shortstack{{{len(shortstack_segments)-1}}}"
        cleaned_inner = cleaned_inner[:idx] + replacement + cleaned_inner[brace_start + len(body) + 2:]
        pos = idx + len(replacement)

    if p_only:
        # Split rows by horizontal rules; keep \\ as in-cell line breaks
        raw_rows = re.split(r"\\(?:toprule|midrule|bottomrule|hline)\b", cleaned_inner)
    else:
        # render_list may have replaced row separators with \x00NL\x00
        cleaned_inner = cleaned_inner.replace("\x00NL\x00", "\\\\")
        raw_rows = re.split(r"\\\\", cleaned_inner)

    def _restore_shortstack(m):
        idx = int(m.group(1))
        body = shortstack_segments[idx]
        body = body.replace("\x00SJ\x00", "\\\\")
        return f"\\shortstack{{{body}}}"
    raw_rows = [re.sub(r"\\shortstack\{(\d+)\}", _restore_shortstack, r) for r in raw_rows]

    rows_data = []
    is_header_row = []
    
    for raw in raw_rows:
        raw = raw.strip()
        # Remove optional spacing arguments from \\[...] that ended up at row start
        raw = re.sub(r"^\[[^\]]*\]\s*", "", raw).strip()
        if not raw or raw.startswith("%"):
            continue
        # Strip table rules that may be embedded at the start of a row (e.g. \midrule\nBloque)
        raw = re.sub(r"\\(toprule|midrule|bottomrule|hline)\b\s*", "", raw).strip()
        raw = re.sub(r"\\addlinespace\b\s*", "", raw).strip()
        if not raw:
            continue
        if _is_rule(raw):
            continue
        # Skip rows that are only layout commands (e.g. \centering before table rows)
        if re.match(r"^\s*\\(?:centering|noindent|par\b|medskip|bigskip|smallskip)\s*$", raw):
            continue
            
        is_hdr = bool(re.match(r"\s*\\rowcolor\{blue!20\}", raw))
        raw = re.sub(r"\\rowcolor\{[^}]*\}", "", raw).strip()
        if not raw:
            continue

        raw = _expand_multicolumn(raw)
        cells = [c.strip() for c in raw.split("&")]
        cells = [c for c in cells if not re.match(r"^p\{", c.strip())]
        # Remove only trailing empty cells; keep leading/middle empty cells
        # because they are needed for \multirow alignment
        while cells and not cells[-1].strip():
            cells.pop()
        
        # Filtrar filas de continuación (longtable)
        if cells and is_continuation_row_cells(cells):
            continue
        
        # Filtrar duplicados exactos
        if cells and is_duplicate_row(cells, [tuple(r) for r in rows_data]):
            continue
        
        if cells:
            rows_data.append(cells)
            is_header_row.append(is_hdr)

    if not rows_data:
        return

    ncols = ncols_hint or max(len(r) for r in rows_data)
    
    for i, row in enumerate(rows_data):
        if len(row) < ncols:
            row.extend([""] * (ncols - len(row)))
        elif len(row) > ncols:
            rows_data[i] = row[:ncols]

    if not col_widths_dxa or len(col_widths_dxa) != ncols:
        col_widths_dxa = _parse_col_widths_dxa("", ncols)

    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row_idx = 0

    for ri, cells in enumerate(rows_data):
        tr = table.rows[ri]
        for ci in range(ncols):
            cell = tr.cells[ci]
            cell_text = cells[ci] if ci < len(cells) else ""
            cell_text = re.sub(r"\\label\{[^}]*\}", "", cell_text)
            cell_text = re.sub(r"\\hspace\*?\{[^}]*\}", "", cell_text)

            # For vertical p-only tables, treat \\ inside a cell as line breaks
            if p_only:
                cell_lines = [ln.strip() for ln in re.split(r"\\\\(?:\[[^\]]*\])?", cell_text) if ln.strip()]
            elif r"\shortstack" in cell_text:
                # Expandir \shortstack{...} en líneas separadas (saltos de línea de celda).
                # Usamos extracción balanceada para soportar \textbf{...} dentro.
                idx = cell_text.find("\\shortstack{")
                brace_start = idx + len("\\shortstack{") - 1
                body = _extract_balanced_braces(cell_text, brace_start)
                if body:
                    cell_lines = [ln.strip() for ln in body.split("\\\\") if ln.strip()]
                else:
                    cell_lines = [cell_text]
            else:
                cell_lines = [cell_text]

            for line_idx, line in enumerate(cell_lines):
                if line_idx == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if (ri == header_row_idx) else WD_ALIGN_PARAGRAPH.LEFT

                # En tablas array (\begin{array} en modo math), envolver cada
                # celda en $...$ para que Pandoc la convierta a OMML, salvo que
                # ya contenga math inline.
                if is_array and "$" not in line:
                    line = f"${line}$"

                # Sanitize any stray marker characters that leaked through
                line = line.replace("\x00SJ\x00", "\\\\").replace("\x00NL\x00", "\\\\")
                if ri == header_row_idx:
                    if '$' in line:
                        parse_inline(p, line, base_sz=PT_SMALL)
                        for run in p.runs:
                            run.bold = True
                    else:
                        _run(p, strip_fmt(line), bold=True, size=PT_SMALL)
                else:
                    parse_inline(p, line, base_sz=PT_SMALL)

    apply_booktabs_style(table)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# List renderer
# ─────────────────────────────────────────────────────────────────────────────

def _get_list_number_abstract_num_id(doc: Document):
    """Find the abstractNumId used by the 'List Number' style in this document."""
    try:
        # 1. Try to find an existing paragraph with List Number and read its numId
        list_number_num_id = None
        for para in doc.paragraphs:
            if para.style and para.style.name == "List Number":
                pPr = para._p.find(qn('w:pPr'))
                if pPr is not None:
                    numPr = pPr.find(qn('w:numPr'))
                    if numPr is not None:
                        numId_el = numPr.find(qn('w:numId'))
                        if numId_el is not None:
                            list_number_num_id = int(numId_el.get(qn('w:val')))
                            break

        # 2. If no paragraph exists yet, create a temp one to get the numId
        if list_number_num_id is None:
            temp = doc.add_paragraph(style="List Number")
            pPr = temp._p.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    numId_el = numPr.find(qn('w:numId'))
                    if numId_el is not None:
                        list_number_num_id = int(numId_el.get(qn('w:val')))
            # Remove temp paragraph
            temp._element.getparent().remove(temp._element)

        # 3. Look up the abstractNumId for this numId in numbering.xml
        if list_number_num_id is not None:
            numbering = doc.part.numbering_part.numbering_definitions._numbering
            for num in numbering.findall(qn('w:num')):
                if int(num.get(qn('w:numId'))) == list_number_num_id:
                    abstractNumId_el = num.find(qn('w:abstractNumId'))
                    if abstractNumId_el is not None:
                        return int(abstractNumId_el.get(qn('w:val')))
        return None
    except Exception:
        return None


def _split_items(inner: str) -> list[str]:
    r"""Split list inner content by \item only at the current nesting level.

    This avoids treating \item commands inside nested itemize/enumerate
    environments as top-level item separators.
    """
    items = []
    current = []
    depth = 0
    i = 0
    n = len(inner)
    while i < n:
        m_begin = re.match(r"\\begin\{(itemize|enumerate)\*?\}", inner[i:])
        if m_begin:
            depth += 1
            current.append(m_begin.group(0))
            i += m_begin.end()
            continue
        m_end = re.match(r"\\end\{(itemize|enumerate)\*?\}", inner[i:])
        if m_end:
            depth = max(0, depth - 1)
            current.append(m_end.group(0))
            i += m_end.end()
            continue
        if depth == 0 and inner.startswith(r"\item", i):
            items.append("".join(current))
            current = []
            i += len(r"\item")
            continue
        current.append(inner[i])
        i += 1
    if current:
        items.append("".join(current))
    # Drop leading empty chunk before the first \item
    if items and not items[0].strip():
        items = items[1:]
    return items


def render_list(doc: Document, inner: str, ordered: bool = False, depth: int = 0, base_dir: Path = Path("."), label_format: str = None):
    r"""
    Recursively parse itemize / enumerate and add list paragraphs.
    Handles nested lists, custom \item[label] labels,
    and block environments (figure, table, center, etc.) inside items.
    """
    bullet_style  = "List Bullet"  if not ordered else "List Number"
    # Try to use indented styles for nesting
    if depth > 0:
        bullet_style = "List Bullet 2" if not ordered else "List Number 2"

    # Split on \item boundaries, respecting nested itemize/enumerate
    items_raw = _split_items(inner)
    
    item_index = 0  # Counter for auto-numbering when no custom label

    # Environments that can appear inside an \item and need special handling
    NESTED_ENV_RE = re.compile(
        r"\\begin\{"
        r"(itemize|enumerate|figure|table\*?|longtable|center|tikzpicture"
        r"|equation\*?|align\*?|gather\*?|multline\*?|eqnarray\*?"
        r"|small|tcolorbox|minipage|tabularx?)"
        r"\*?\}"
    )

    # Sectioning commands that can appear inside an \item (e.g., \subsubsection)
    SECTION_RE = re.compile(r"\\(section|subsection|subsubsection|paragraph)(\*)?\s*\{")

    for item_raw in items_raw:
        item_raw = item_raw.strip()
        # Preserve explicit line breaks as paragraph separators
        item_raw = item_raw.replace("\\\\", "\x00NL\x00")
        item_raw = item_raw.replace("\\newline", "\x00NL\x00")
        # Normalize internal newlines/spaces from multi-line LaTeX source
        item_raw = re.sub(r"\s+", " ", item_raw)
        if not item_raw:
            continue

        # Custom label:  [label] text
        label_m = re.match(r"^\[([^\]]*)\]\s*", item_raw)
        custom_label = None
        has_custom_label = False
        is_enumitem_label = False
        if label_m:
            custom_label = _clean(label_m.group(1))
            item_raw     = item_raw[label_m.end():]
            has_custom_label = True
        else:
            item_index += 1  # Only increment for auto-numbered items
            if ordered and label_format:
                custom_label = _format_enumitem_label(label_format, item_index)
                has_custom_label = True
                is_enumitem_label = True

        # When custom label exists, use normal paragraph (not list style) to avoid double numbering
        # This also makes text use 'Normal' style (justified, more width)
        if has_custom_label:
            para_style = "Normal"  # Use Normal style for custom labels (justified, more width)
        else:
            para_style = bullet_style

        # Helper to render a paragraph of item text
        def _render_item_text(text: str, use_label: bool, continuation: bool = False):
            if not text and not use_label:
                return
            # Split explicit line breaks into separate paragraphs
            parts = [p.strip() for p in text.split("\x00NL\x00")]
            parts = [p for p in parts if p]
            if not parts and not use_label:
                return

            first = True
            for part in (parts if parts else [""]):
                if ordered and not has_custom_label and not use_label and not continuation and first:
                    # Manual numbering ensures each enumerate restarts at 1
                    p = doc.add_paragraph(style="Normal")
                    p.paragraph_format.left_indent = Inches(0.5 * (depth + 1))
                    p.paragraph_format.first_line_indent = Inches(-0.25)
                    run = p.add_run(f"{item_index}.  ")
                    run.font.name = FONT
                    run.font.size = Pt(PT_NORM)
                    run.font.color.rgb = C_BLACK
                    parse_inline(p, part)
                    _maybe_add_tab_stop(p, part)
                else:
                    if continuation or not first:
                        # Continuation paragraph: normal style, same indentation as list items, no number
                        p = doc.add_paragraph(style="Normal")
                        p.paragraph_format.left_indent = Inches(0.5 * (depth + 1))
                    else:
                        p = doc.add_paragraph(style=para_style)
                        if has_custom_label:
                            if is_enumitem_label:
                                p.paragraph_format.left_indent = Inches(0.5 * (depth + 1))
                                p.paragraph_format.first_line_indent = Inches(-0.35)
                            else:
                                p.paragraph_format.left_indent = Inches(0.3 * (depth + 1))
                    if first and use_label and custom_label:
                        n_runs_before = len(p.runs)
                        parse_inline(p, custom_label)
                        for run in p.runs[n_runs_before:]:
                            # enumitem default labels should match LaTeX formatting (not bold)
                            if not is_enumitem_label:
                                run.bold = True
                            run.font.name = FONT
                        _run(p, "  ", size=PT_NORM)
                    parse_inline(p, part)
                    _maybe_add_tab_stop(p, part)
                first = False

        # Helper to render a sectioning command found inside an item
        def _render_section_heading(cmd: str, title: str):
            level_map = {
                "section": 2,
                "subsection": 3,
                "subsubsection": 4,
                "paragraph": 5,
            }
            size_map = {2: 14, 3: 13, 4: 12, 5: 11}
            level = level_map.get(cmd, 4)
            style_name = f"Heading {level}"
            p = doc.add_paragraph(style=style_name)
            run = p.add_run(strip_fmt(title))
            run.font.name = FONT
            run.font.size = Pt(size_map.get(level, 12))
            run.font.bold = True
            run.font.color.rgb = C_BLACK

        # Process the item extracting nested environments and sectioning commands in order
        remaining = item_raw
        label_used = False
        is_continuation = False  # True for text that follows a nested environment

        while remaining:
            env_match = NESTED_ENV_RE.search(remaining)
            section_match = SECTION_RE.search(remaining)

            # Choose the nearest structural marker
            if env_match and section_match:
                if env_match.start() < section_match.start():
                    section_match = None
                else:
                    env_match = None

            if not env_match and not section_match:
                _render_item_text(remaining, use_label=not label_used, continuation=is_continuation)
                break

            if section_match:
                pre_text = remaining[:section_match.start()].strip()
                cmd = section_match.group(1)
                brace_start = section_match.start() + section_match.end() - section_match.start() - 1
                # section_match.end() points after '{', so brace_start is the position of '{'
                brace_start = section_match.end() - 1
                title = _extract_balanced_braces(remaining, brace_start)
                if title is None:
                    _render_item_text(remaining, use_label=not label_used, continuation=is_continuation)
                    break
                end_pos = brace_start + len(title) + 2
                post_text = remaining[end_pos:].strip()

                _render_item_text(pre_text, use_label=not label_used, continuation=is_continuation)
                label_used = True
                _render_section_heading(cmd, title)
                remaining = post_text
                is_continuation = True
                continue

            pre_text = remaining[:env_match.start()].strip()
            env_name = env_match.group(1)

            result = extract_env(remaining, env_name, env_match.start())
            if not result:
                _render_item_text(remaining, use_label=not label_used, continuation=is_continuation)
                break

            _, end_pos, env_inner = result
            env_text = remaining[env_match.start():end_pos]
            post_text = remaining[end_pos:].strip()

            # Render text before the environment
            _render_item_text(pre_text, use_label=not label_used, continuation=is_continuation)
            label_used = True

            # Render the environment itself
            if env_name in ("itemize", "enumerate"):
                opt_arg, env_inner_clean = _extract_optional_arg(env_inner)
                nested_label = _parse_enumitem_label(opt_arg) if opt_arg else None
                # Apply enumitem default label for nested enumerate if no explicit option
                if env_name == "enumerate" and nested_label is None:
                    nested_label = _get_enumitem_default(depth + 2)
                render_list(doc, env_inner_clean, ordered=(env_name == "enumerate"), depth=depth + 1, base_dir=base_dir, label_format=nested_label)
            else:
                _walk(doc, env_text, base_dir)

            # Continue scanning the remainder of the item for more nested environments
            remaining = post_text
            is_continuation = True

        # If item was empty but had a label, make sure at least the label is rendered
        if not remaining and not label_used and custom_label:
            _render_item_text("", use_label=True)
def parse_body(doc: Document, source: str, base_dir: Path):
    """Walk the entire source and render into `doc`."""
    global PREAMBLE_GLOBAL
    
    # Reiniciar contadores
    TABLE_COUNTER[0] = 0
    TIKZ_COUNTER[0] = 0
    FIGURE_COUNTER[0] = 0

    # Guardar preámbulo para compilar tikzpictures
    preamble_end = source.find("\\begin{document}")
    PREAMBLE_GLOBAL = source[:preamble_end] if preamble_end != -1 else ""

    # Extraer rutas de imágenes del preámbulo
    global GRAPHICSPATHS
    GRAPHICSPATHS = _extract_graphicspaths(PREAMBLE_GLOBAL)

    # Extraer opciones por defecto de enumitem del preámbulo
    global ENUMITEM_DEFAULTS
    ENUMITEM_DEFAULTS = _extract_enumitem_defaults(PREAMBLE_GLOBAL)

    # ── remove comments ──────────────────────────────────────────────────────
    # Elimina comentarios (% hasta fin de línea) pero preserva \% (porcentaje literal)
    source = re.sub(r"(?m)(?<!\\)%.*$", "", source)

    # ── title block ──────────────────────────────────────────────────────────
    def _extract_one(cmd, text):
        # Find command \cmd not followed by another letter
        m = re.search(r"\\" + cmd + r"(?![a-zA-Z])", text)
        if not m:
            return ""
        brace_idx = text.find("{", m.end())
        if brace_idx == -1:
            return ""
        content = _extract_balanced_braces(text, brace_idx)
        return content if content is not None else ""

    title_raw  = _extract_one("title",  source)
    author_raw = _extract_one("author", source)
    date_raw   = _extract_one("date",   source)

    # Extract optional title image (\includegraphics inside \title)
    title_img_path = None
    title_img_width = None
    title_img_height = None
    img_m = re.search(r"\\includegraphics(?:\[([^\]]*)\])?\{([^}]*)\}", title_raw)
    if img_m:
        opts = img_m.group(1) or ""
        img_name = img_m.group(2)
        title_img_path = resolve_image_path(img_name, base_dir)
        if title_img_path and title_img_path.exists():
            title_img_width, title_img_height = _parse_graphics_size(opts, title_img_path)
        # Remove the \includegraphics from title_raw for text extraction
        title_raw = title_raw[:img_m.start()] + title_raw[img_m.end():]

    # Use the full cleaned title text (no hardcoded defaults)
    title_text = strip_fmt(title_raw).strip()
    # If title is empty after cleaning, try \textbf as fallback
    if not title_text:
        tb_m = re.search(r"\\textbf\{([^}]+)\}", title_raw)
        if tb_m:
            title_text = strip_fmt(tb_m.group(1))

    author_text = strip_fmt(author_raw)
    date_text   = strip_fmt(date_raw)

    has_title_content = bool(title_text or author_text or date_text or (title_img_path and title_img_path.exists()))
    if has_title_content:
        add_title_page(doc, title_img_path, title_img_width, title_img_height, title_text, author_text, date_text)
        _safe_page_break(doc)

    # ── Table of Contents ────────────────────────────────────────────────────
    add_toc(doc)
    _safe_page_break(doc)
    
    # ── isolate the document body ─────────────────────────────────────────────
    begin_doc = re.search(r"\\begin\{document\}", source)
    end_doc   = re.search(r"\\end\{document\}",   source)
    if not begin_doc:
        return
    body = source[begin_doc.end(): end_doc.start() if end_doc else len(source)]

    # Remove \title, \author, \date blocks from body so _walk doesn't re-render them
    def _remove_cmd_block(cmd: str, text: str) -> str:
        m = re.search(r"\\" + cmd + r"(?![a-zA-Z])", text)
        if not m:
            return text
        brace_idx = text.find("{", m.end())
        if brace_idx == -1:
            return text
        content = _extract_balanced_braces(text, brace_idx)
        if content is None:
            return text
        end_idx = brace_idx + 1 + len(content) + 1
        start_idx = m.start()
        return text[:start_idx] + text[end_idx:]

    for cmd in ("title", "author", "date"):
        body = _remove_cmd_block(cmd, body)

    # ── resolve \\ref{...} to numbers ────────────────────────────────────────
    labels = pre_scan_labels(body)
    body = resolve_refs(body, labels)

    # ── strip fancyhdr commands that leak into body text ──────────────────────
    body = strip_fancyhdr_cmds(body)

    # ── scan bibliography ─────────────────────────────────────────────────────
    bib_map, bib_inner, body, bib_title = pre_scan_bibliography(body)

    # ── resolve \\cite{...} to numbers ───────────────────────────────────────
    body = resolve_cites(body, bib_map)

    # ── scan figures and tables for lists ────────────────────────────────────
    figures = pre_scan_figures(body)
    tables = pre_scan_tables(body)

    # ── walk the body ─────────────────────────────────────────────────────────
    _walk(doc, body, base_dir, figures, tables)

    # ── render bibliography ───────────────────────────────────────────────────
    if bib_inner is not None:
        _safe_page_break(doc)
        if bib_title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(bib_title.upper())
            run.font.name = FONT
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = C_BLACK
        _render_bibliography(doc, bib_inner, bib_map)


def strip_fancyhdr_cmds(text: str) -> str:
    """Remove fancyhdr setup commands that leak into the document body."""
    # Single-line fancyhdr / header-footer commands
    text = re.sub(
        r"(?m)^\s*\\(?:pagestyle|thispagestyle|fancyhf|rhead|lhead|chead|fancyfoot|rfoot|lfoot|cfoot|headrulewidth|footrulewidth)\b.*(?:\n|\r\n?)",
        "", text,
    )
    # \renewcommand{\headrulewidth}{...} or {\footrulewidth}{...}
    text = re.sub(
        r"(?m)^\s*\\renewcommand\s*\*?\s*\{\s*\\(?:headrulewidth|footrulewidth)\s*\}\s*\{[^}]*\}\s*(?:\n|\r\n?)",
        "", text,
    )
    # \setlength\headheight{...}
    text = re.sub(
        r"(?m)^\s*\\setlength\s*\\headheight\s*\{[^}]*\}\s*(?:\n|\r\n?)",
        "", text,
    )
    return text


# Significant structural tokens
_STRUCT = re.compile(
    r"\\(?:chapter|section|subsection|subsubsection|paragraph|begin|end|newpage|clearpage)\*?\b"
    r"|\\appendix\b"
    r"|\\maketitle\b|\\tableofcontents\b|\\listoffigures\b|\\listoftables\b"
    r"|\\includegraphics\b"
)

# Display math
_DISP_MATH = re.compile(r"\\\[(.*?)\\\]", re.DOTALL)


def pre_scan_labels(text: str) -> dict:
    """Scan text to map \\label{...} to figure/table numbers."""
    labels = {}
    fig_counter = 0
    table_counter = 0

    # Figures (including figure*)
    fig_pattern = re.compile(r"\\begin\{figure\*?\}")
    pos = 0
    while True:
        m = fig_pattern.search(text, pos)
        if not m:
            break
        idx = m.start()
        result = extract_env(text, "figure", idx)
        if result:
            _, end_pos, inner = result
            if "\\caption{" in inner:
                fig_counter += 1
                for label_m in re.finditer(r"\\label\{([^}]*)\}", inner):
                    labels[label_m.group(1)] = str(fig_counter)
            pos = end_pos
        else:
            pos = idx + 1

    # Tables
    for env_name in ("table", "table*", "longtable"):
        pos = 0
        search_str = f"\\begin{{{env_name}}}"
        while True:
            idx = text.find(search_str, pos)
            if idx == -1:
                break
            result = extract_env(text, env_name, idx)
            if result:
                _, end_pos, inner = result
                if "\\caption{" in inner:
                    table_counter += 1
                    for label_m in re.finditer(r"\\label\{([^}]*)\}", inner):
                        labels[label_m.group(1)] = str(table_counter)
                pos = end_pos
            else:
                pos = idx + 1

    # Equations (all types, in order of appearance)
    eq_counter = 0
    pos = 0
    eq_envs = ("equation", "equation*", "align", "align*", "gather", "gather*", "multline", "multline*")
    while pos < len(text):
        next_idx = len(text)
        is_env = False
        env_name = None

        for en in eq_envs:
            idx = text.find(f"\\begin{{{en}}}", pos)
            if idx != -1 and idx < next_idx:
                next_idx = idx
                is_env = True
                env_name = en

        # Use _DISP_MATH regex to find display math (ignores \\[ like in \\[\bigskipamount])
        dm_match = _DISP_MATH.search(text, pos)
        idx_dm = dm_match.start() if dm_match else -1
        if idx_dm != -1 and idx_dm < next_idx:
            next_idx = idx_dm
            is_env = False

        if next_idx == len(text):
            break

        if is_env:
            result = extract_env(text, env_name, next_idx)
            if result:
                _, end_pos, inner = result
                # align/gather/etc. may contain multiple equations (one per \\ line);
                # number each labeled line sequentially.
                lines = re.split(r"\\\\(?:\[[^\]]*\])?", inner)
                for line in lines:
                    label_m = re.search(r"\\label\{([^}]*)\}", line)
                    if label_m:
                        eq_counter += 1
                        labels[label_m.group(1)] = str(eq_counter)
                pos = end_pos
            else:
                pos = next_idx + 1
        else:
            dm = _DISP_MATH.match(text, next_idx)
            if not dm:
                break
            block = dm.group(0)
            label_m = re.search(r"\\label\{([^}]*)\}", block)
            if label_m:
                eq_counter += 1
                labels[label_m.group(1)] = str(eq_counter)
            pos = dm.end()

    # Section / chapter labels
    sec_counters = {"chapter": 0, "section": 0, "subsection": 0, "subsubsection": 0, "paragraph": 0}
    has_chapters = bool(re.search(r"\\chapter\b", text))
    appendix_pos = text.find("\\appendix")

    sec_matches = list(re.finditer(
        r"\\(chapter|section|subsection|subsubsection|paragraph)\*?\s*\{[^{}]*\}",
        text
    ))

    for i, m in enumerate(sec_matches):
        cmd = m.group(1)
        start_pos = m.end()
        end_pos = sec_matches[i + 1].start() if i + 1 < len(sec_matches) else len(text)

        if cmd == "chapter":
            sec_counters["chapter"] += 1
            sec_counters["section"] = 0
            sec_counters["subsection"] = 0
            sec_counters["subsubsection"] = 0
            sec_counters["paragraph"] = 0
        elif cmd == "section":
            sec_counters["section"] += 1
            sec_counters["subsection"] = 0
            sec_counters["subsubsection"] = 0
            sec_counters["paragraph"] = 0
        elif cmd == "subsection":
            sec_counters["subsection"] += 1
            sec_counters["subsubsection"] = 0
            sec_counters["paragraph"] = 0
        elif cmd == "subsubsection":
            sec_counters["subsubsection"] += 1
            sec_counters["paragraph"] = 0
        elif cmd == "paragraph":
            sec_counters["paragraph"] += 1

        appendix_mode = appendix_pos != -1 and m.start() > appendix_pos

        if not has_chapters:
            if cmd == "section":
                num_str = f"{sec_counters['section']}"
            elif cmd == "subsection":
                num_str = f"{sec_counters['section']}.{sec_counters['subsection']}"
            elif cmd == "subsubsection":
                num_str = f"{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}"
            elif cmd == "paragraph":
                num_str = f"{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}.{sec_counters['paragraph']}"
            else:
                num_str = ""
        else:
            ch_label = chr(ord('A') + sec_counters['chapter'] - 1) if appendix_mode else str(sec_counters['chapter'])
            if cmd == "chapter":
                num_str = ch_label
            elif cmd == "section":
                num_str = f"{ch_label}.{sec_counters['section']}"
            elif cmd == "subsection":
                num_str = f"{ch_label}.{sec_counters['section']}.{sec_counters['subsection']}"
            elif cmd == "subsubsection":
                num_str = f"{ch_label}.{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}"
            elif cmd == "paragraph":
                num_str = f"{ch_label}.{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}.{sec_counters['paragraph']}"
            else:
                num_str = ""

        # Look for the first \label within a short window after the section command
        search_end = min(end_pos, start_pos + 300)
        label_m = re.search(r"\\label\{([^}]*)\}", text[start_pos:search_end])
        if label_m and label_m.group(1) not in labels:
            labels[label_m.group(1)] = num_str

    return labels


def pre_scan_figures(text: str) -> list[tuple[int, str]]:
    """Return list of (number, plain_caption) for every figure environment."""
    figures = []
    for env_name in ("figure", "figure*"):
        pos = 0
        search_str = f"\\begin{{{env_name}}}"
        while True:
            idx = text.find(search_str, pos)
            if idx == -1:
                break
            result = extract_env(text, env_name, idx)
            if result:
                _, end_pos, inner = result
                # Remove subfigure environments so we only keep the main caption
                inner_no_sub = inner
                sub_pos = 0
                while True:
                    sub_idx = inner_no_sub.find("\\begin{subfigure}", sub_pos)
                    if sub_idx == -1:
                        break
                    sub_res = extract_env(inner_no_sub, "subfigure", sub_idx)
                    if sub_res:
                        _, sub_end, _ = sub_res
                        inner_no_sub = inner_no_sub[:sub_idx] + inner_no_sub[sub_end:]
                    else:
                        sub_pos = sub_idx + 1
                cap = _extract_caption_content(inner_no_sub)
                if cap is not None:
                    cap_clean = strip_fmt(cap)
                    cap_clean = re.sub(r"\\label\{[^}]*\}", "", cap_clean).strip()
                    if cap_clean:
                        figures.append((len(figures) + 1, cap_clean))
                pos = end_pos
            else:
                pos = idx + 1
    return figures


def pre_scan_tables(text: str) -> list[tuple[int, str]]:
    """Return list of (number, plain_caption) for every table environment."""
    tables = []
    for env_name in ("table", "table*", "longtable"):
        pos = 0
        search_str = f"\\begin{{{env_name}}}"
        while True:
            idx = text.find(search_str, pos)
            if idx == -1:
                break
            result = extract_env(text, env_name, idx)
            if result:
                _, end_pos, inner = result
                cap = _extract_caption_content(inner)
                if cap is not None:
                    cap_clean = strip_fmt(cap)
                    cap_clean = re.sub(r"\\label\{[^}]*\}", "", cap_clean).strip()
                    if cap_clean:
                        tables.append((len(tables) + 1, cap_clean))
                pos = end_pos
            else:
                pos = idx + 1
    return tables


def resolve_refs(text: str, labels: dict) -> str:
    """Replace \\ref{...} and \\eqref{...} with the corresponding number from labels."""
    def repl(m):
        label_name = m.group(1)
        return labels.get(label_name, f"[ref:{label_name}]")
    text = re.sub(r"\\eqref\{([^}]*)\}", repl, text)
    text = re.sub(r"\\ref\{([^}]*)\}", repl, text)
    return text


def pre_scan_bibliography(text: str) -> tuple[dict, str | None, str, str]:
    """
    Extract \\begin{thebibliography}...\\end{thebibliography} from text.
    Returns (bib_map, bib_inner, text_without_bibliography, bib_title).
    bib_map: {cite_key: number}
    """
    pattern = r"\\begin\{thebibliography\}\{[^}]*\}(.*?)\\end\{thebibliography\}"
    m = re.search(pattern, text, re.DOTALL)
    if not m:
        return {}, None, text, ""

    bib_inner = m.group(1)
    text_without = text[:m.start()] + text[m.end():]

    bib_map = {}
    counter = 1
    for bm in re.finditer(r"\\bibitem\{([^}]+)\}", bib_inner):
        key = bm.group(1).strip()
        if key and key not in bib_map:
            bib_map[key] = counter
            counter += 1

    # Look for \\renewcommand{\\refname}{...} or \\renewcommand{\\bibname}{...}
    # in the text before thebibliography
    prefix = text[max(0, m.start()-500):m.start()]
    title_m = re.search(r"\\renewcommand\{\\(?:refname|bibname)\}\{((?:[^{}]|\{[^{}]*\})*)\}", prefix)
    bib_title = title_m.group(1) if title_m else CONFIG.bib_title_default

    # Remove the \\renewcommand from text_without so it doesn't render as stray text
    if title_m:
        rc_pattern = r"\\renewcommand\{\\(?:refname|bibname)\}\{" + re.escape(title_m.group(1)) + r"\}"
        text_without = re.sub(rc_pattern, "", text_without)

    return bib_map, bib_inner, text_without, bib_title


def resolve_cites(text: str, bib_map: dict) -> str:
    """Replace \\cite{key1,key2} with [n1,n2] using bib_map."""
    def repl(m):
        keys = [k.strip() for k in m.group(1).split(",")]
        nums = [str(bib_map.get(k, k)) for k in keys if k]
        return "[" + ", ".join(nums) + "]" if nums else ""

    text = re.sub(r"\\cite\{([^}]+)\}", repl, text)
    return text


def _render_bibliography(doc: Document, bib_inner: str, bib_map: dict):
    """
    Render thebibliography inner content as a numbered list of paragraphs.
    Each \\bibitem{key} becomes a paragraph with a hanging indent.
    """
    # Split by \bibitem{...}
    items = re.split(r"\\bibitem\{([^}]+)\}", bib_inner)
    if len(items) < 2:
        return

    i = 1
    while i < len(items):
        key = items[i].strip()
        text = items[i + 1] if i + 1 < len(items) else ""
        num = bib_map.get(key, "")
        if text:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            _run(p, f"[{num}] ", size=PT_NORM)
            # Strip leading whitespace/newlines
            text = text.strip()
            if text:
                parse_inline(p, text, base_sz=PT_NORM)
        i += 2


def _insert_list_of_figures(doc: Document, figures: list[tuple[int, str]],
                            config: Config | None = None):
    """Insert a real Word 'List of Figures' TOC field."""
    if config is None:
        config = CONFIG
    if not figures:
        return
    _insert_toc_field(
        doc,
        r'TOC \f F \h',
        config.lof_title,
        config.lof_placeholder
    )


def _insert_list_of_tables(doc: Document, tables: list[tuple[int, str]],
                           config: Config | None = None):
    """Insert a real Word 'List of Tables' TOC field."""
    if config is None:
        config = CONFIG
    if not tables:
        return
    _insert_toc_field(
        doc,
        r'TOC \f T \h',
        config.lot_title,
        config.lot_placeholder
    )


def _walk(doc: Document, text: str, base_dir: Path, figures: list[tuple[int, str]] | None = None, tables: list[tuple[int, str]] | None = None):
    """
    Scan `text` linearly.  At each structural command do the right thing,
    otherwise accumulate plain/inline text and flush as a paragraph.
    """
    pending: list[str] = []

    # Section / chapter counters
    sec_counters = {"chapter": 0, "section": 0, "subsection": 0, "subsubsection": 0, "paragraph": 0}
    appendix_mode = False
    has_chapters = False

    # Current font state propagated across paragraph breaks
    fmt_bold = False
    fmt_italic = False
    fmt_size = PT_NORM

    def flush():
        nonlocal pending, fmt_bold, fmt_italic, fmt_size
        raw = " ".join(pending).strip()
        if raw:
            _add_para(doc, raw, init_bold=fmt_bold, init_italic=fmt_italic, init_size=fmt_size)
            # Propagate bold/italic, but reset size between paragraphs.
            # LaTeX size switches (\small, \footnotesize, ...) are scoped to
            # their paragraph/group and must not leak into the rest of the document.
            fmt_bold, fmt_italic, _ = _scan_format_switches(raw, fmt_bold, fmt_italic, PT_NORM)
        pending = []
    
    def format_section_number(cmd: str) -> str:
        """Generate section number like 1, 1.1, 1.1.1, or A.1"""
        ch = sec_counters['chapter']
        if not has_chapters:
            # No chapters → flat numbering as before
            if cmd == "section":
                return f"{sec_counters['section']}"
            elif cmd == "subsection":
                return f"{sec_counters['section']}.{sec_counters['subsection']}"
            elif cmd == "subsubsection":
                return f"{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}"
            elif cmd == "paragraph":
                return f"{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}.{sec_counters['paragraph']}"
            return ""
        # With chapters → hierarchical numbering
        ch_label = chr(ord('A') + ch - 1) if appendix_mode else str(ch)
        if cmd == "chapter":
            return ch_label
        elif cmd == "section":
            return f"{ch_label}.{sec_counters['section']}"
        elif cmd == "subsection":
            return f"{ch_label}.{sec_counters['section']}.{sec_counters['subsection']}"
        elif cmd == "subsubsection":
            return f"{ch_label}.{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}"
        elif cmd == "paragraph":
            return f"{ch_label}.{sec_counters['section']}.{sec_counters['subsection']}.{sec_counters['subsubsection']}.{sec_counters['paragraph']}"
        return ""

    pos = 0
    n   = len(text)

    while pos < n:
        # Display math
        dm = _DISP_MATH.match(text, pos)
        if dm:
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            omml = None
            if TEMP_DIR and PANDOC_PATH:
                try:
                    omml = latex_to_omml_element(dm.group(0), TEMP_DIR, PANDOC_PATH)
                except Exception:
                    pass
            if omml is not None:
                p._p.append(omml)
            else:
                math_text = strip_fmt(dm.group(1))
                _run(p, math_text, italic=True, size=PT_NORM)
            pos = dm.end()
            continue

        # Page break
        pb = re.match(r"\\(?:newpage|clearpage)\b", text[pos:])
        if pb:
            flush()
            _safe_page_break(doc)
            pos += pb.end()
            continue

        # Skip \maketitle, \tableofcontents (already handled)
        skip = re.match(r"\\(?:maketitle|tableofcontents)\b", text[pos:])
        if skip:
            pos += skip.end()
            continue
        
        # \listoffigures
        lof_m = re.match(r"\\listoffigures\b", text[pos:])
        if lof_m:
            flush()
            if figures:
                _insert_list_of_figures(doc, figures)
            pos += lof_m.end()
            continue
        
        # \listoftables
        lot_m = re.match(r"\\listoftables\b", text[pos:])
        if lot_m:
            flush()
            if tables:
                _insert_list_of_tables(doc, tables)
            pos += lot_m.end()
            continue
        
        # \appendix — switch to appendix mode
        app_m = re.match(r"\\appendix\b", text[pos:])
        if app_m:
            appendix_mode = True
            pos += app_m.end()
            continue
        
        # Chapter headings
        ch_m = re.match(
            r"\\(chapter)\*?\s*\{((?:[^{}]|\{[^{}]*\})*)\}",
            text[pos:]
        )
        if ch_m:
            flush()
            _safe_page_break(doc)
            has_chapters = True
            sec_counters["chapter"] += 1
            sec_counters["section"] = 0
            sec_counters["subsection"] = 0
            sec_counters["subsubsection"] = 0
            sec_counters["paragraph"] = 0
            htxt = strip_fmt(ch_m.group(2))
            num_str = format_section_number("chapter")
            prefix = CONFIG.appendix_prefix if appendix_mode else CONFIG.chapter_prefix
            full_title = f"{prefix}{num_str}  {htxt}"
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(full_title)
            run.font.name = FONT
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.all_caps = True
            run.font.color.rgb = C_BLACK
            pos += ch_m.end()
            continue

        # Section headings
        sec_m = re.match(
            r"\\(section|subsection|subsubsection|paragraph)(\*)?\s*\{((?:[^{}]|\{[^{}]*\})*)\}",
            text[pos:]
        )
        if sec_m:
            flush()
            cmd      = sec_m.group(1)
            starred  = sec_m.group(2) is not None
            htxt     = strip_fmt(sec_m.group(3))

            # Starred versions (\section*) are not numbered and do not update counters
            if not starred:
                # Update counters
                if cmd == "section":
                    sec_counters["section"] += 1
                    sec_counters["subsection"] = 0
                    sec_counters["subsubsection"] = 0
                    sec_counters["paragraph"] = 0
                elif cmd == "subsection":
                    sec_counters["subsection"] += 1
                    sec_counters["subsubsection"] = 0
                    sec_counters["paragraph"] = 0
                elif cmd == "subsubsection":
                    sec_counters["subsubsection"] += 1
                    sec_counters["paragraph"] = 0
                elif cmd == "paragraph":
                    sec_counters["paragraph"] += 1

            # Get level and formatted number
            if has_chapters:
                level = {"section": 2, "subsection": 3,
                         "subsubsection": 4, "paragraph": 5}[cmd]
                sizes = [16, 14, 13, 12, 11]
            else:
                level = {"section": 1, "subsection": 2,
                         "subsubsection": 3, "paragraph": 4}[cmd]
                sizes = [16, 14, 12, 11]
            num_str = "" if starred else format_section_number(cmd)

            # Create heading with number - use explicit style for TOC recognition
            full_title = f"{num_str}  {htxt}".strip()
            style_name = f"Heading {level}"
            p = doc.add_paragraph(style=style_name)
            run = p.add_run(full_title)
            run.font.name = FONT
            run.font.size = Pt(sizes[level-1])
            run.font.bold = True
            run.font.color.rgb = C_BLACK
            pos += sec_m.end()
            continue

        # \begin{env}
        env_m = re.match(r"\\begin\{(\w+)\*?\}", text[pos:])
        if env_m:
            env_name = env_m.group(1)
            flush()
            result = extract_env(text, env_name, pos)
            if result is None:
                pos += env_m.end()
                continue
            _, end_pos, inner = result

            if env_name in ("itemize", "enumerate", "description"):
                # Extract optional arguments like [label=..., leftmargin=...]
                opt_arg, inner = _extract_optional_arg(inner)
                label_format = _parse_enumitem_label(opt_arg) if opt_arg else None
                # Apply enumitem default label for enumerate if no explicit option
                if env_name == "enumerate" and label_format is None:
                    label_format = _get_enumitem_default(1)
                render_list(doc, inner, ordered=(env_name == "enumerate"), base_dir=base_dir, label_format=label_format)

            elif env_name in ("table", "table*"):
                cap_txt = _extract_caption_content(inner)
                caption = cap_txt if cap_txt is not None else ""
                rendered = False
                # Find tabular/tabularx inside
                for tenv in ("tabularx", "tabular"):
                    tr = extract_env(inner, tenv)
                    if tr:
                        _, _, tab_inner = tr
                        # remove column spec: first {...}
                        col_spec, tab_inner = _strip_column_spec(tab_inner)
                        # remove \resizebox wrapper
                        tab_inner = re.sub(r"\\resizebox\{[^}]*\}\{[^}]*\}\{?\s*", "", tab_inner)
                        tab_inner = re.sub(r"\s*\}?\s*$", "", tab_inner)
                        render_table(doc, tab_inner, caption, col_spec=col_spec)
                        rendered = True
                        break
                if not rendered:
                    # Tabla definida como \begin{array} dentro de math display
                    arr_match = re.search(r"\$\s*\\begin\{array\}(.*?)\\end\{array\}\s*\$", inner, re.DOTALL)
                    if arr_match:
                        tab_inner = arr_match.group(1)
                        col_spec, tab_inner = _strip_column_spec(tab_inner)
                        render_table(doc, tab_inner, caption, col_spec=col_spec, is_array=True)

            elif env_name == "longtable":
                cap_txt = _extract_caption_content(inner)
                caption = cap_txt if cap_txt is not None else ""
                # strip column spec
                col_spec, inner = _strip_column_spec(inner)
                # strip \endfirsthead...\endhead  and  \endfoot...\endlastfoot
                inner = re.sub(r"\\endfirsthead.*?\\endhead",     "", inner, flags=re.DOTALL)
                inner = re.sub(r"\\endfoot.*?\\endlastfoot",       "", inner, flags=re.DOTALL)
                inner = _remove_captions(inner)
                render_table(doc, inner, caption, col_spec=col_spec)

            elif env_name in ("tabular", "tabularx"):
                # tabular/tabularx directo (sin \begin{table} wrapper)
                col_spec, inner = _strip_column_spec(inner)
                # remove \resizebox wrapper
                inner = re.sub(r"\\resizebox\{[^}]*\}\{[^}]*\}\{?\s*", "", inner)
                inner = re.sub(r"\s*\}?\s*$", "", inner)
                render_table(doc, inner, caption="", col_spec=col_spec)

            elif env_name == "center":
                n_para_before = len(doc.paragraphs)
                n_tables_before = len(doc.tables)
                _walk(doc, inner, base_dir)
                # Center all paragraphs and tables added inside this center environment
                for p in doc.paragraphs[n_para_before:]:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for t in doc.tables[n_tables_before:]:
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER

            elif env_name == "tikzpicture":
                tikz_code = r"\begin{tikzpicture}" + inner + r"\end{tikzpicture}"
                png_path = compile_tikz_to_png(tikz_code, PREAMBLE_GLOBAL, TEMP_DIR, TIKZ_COUNTER[0])
                TIKZ_COUNTER[0] += 1
                if png_path and png_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if not _insert_image_with_size(p, png_path, max_width_inches=6.3):
                        _run(p, "[DIAGRAM]", italic=True, color=C_GRAY)
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _run(p, "[DIAGRAM — see PDF version]", italic=True, color=C_GRAY)

            elif env_name == "figure":
                # Intentar renderizar como tabla si hay múltiples \subfigure
                rendered_as_table = _render_figure_as_table(doc, inner, base_dir)
                
                if not rendered_as_table:
                    # Procesar TODAS las \includegraphics dentro del entorno figure
                    img_matches = list(re.finditer(r"\\includegraphics(?:\[([^\]]*)\])?\{([^}]*)\}", inner))
                    if img_matches:
                        for img_m in img_matches:
                            opts = img_m.group(1) or ""
                            img_name = img_m.group(2)
                            img_path = resolve_image_path(img_name, base_dir)
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if img_path and img_path.exists():
                                try:
                                    from PIL import Image
                                    img = Image.open(str(img_path))
                                    px_w, px_h = img.size
                                    nat_w = px_w / 96.0
                                    nat_h = px_h / 96.0
                                    w, h = _parse_graphics_size(opts, img_path)
                                    # Si solo especificaron height, calcular width proporcional
                                    if h is not None and w is None:
                                        w = nat_w * (h / nat_h)
                                    # Si no se especificó nada, usar ancho por defecto
                                    if w is None:
                                        w = 4.0
                                        h = nat_h * (w / nat_w)
                                    # Si excede el ancho de página, escalar proporcionalmente
                                    if w > TEXT_WIDTH_INCHES:
                                        scale = TEXT_WIDTH_INCHES / w
                                        w *= scale
                                        h *= scale
                                    p.add_run().add_picture(str(img_path), width=Inches(w), height=Inches(h))
                                except Exception:
                                    _run(p, f"[Figure: {img_name}]", italic=True)
                            else:
                                _run(p, f"[Figure: {img_name}]", italic=True)
                    else:
                        # Check for tikzpicture inside figure
                        tikz_result = extract_env(inner, "tikzpicture")
                        if tikz_result:
                            _, _, tikz_inner = tikz_result
                            tikz_code = r"\begin{tikzpicture}" + tikz_inner + r"\end{tikzpicture}"
                            png_path = compile_tikz_to_png(tikz_code, PREAMBLE_GLOBAL, TEMP_DIR, TIKZ_COUNTER[0])
                            TIKZ_COUNTER[0] += 1
                            if png_path and png_path.exists():
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                if not _insert_image_with_size(p, png_path, max_width_inches=6.3):
                                    _run(p, "[DIAGRAM]", italic=True, color=C_GRAY)
                            else:
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                _run(p, "[DIAGRAM — see PDF version]", italic=True, color=C_GRAY)
                
                # Caption principal del figure (siempre, independientemente del modo de renderizado)
                fig_cap_txt = _extract_caption_content(inner)
                if fig_cap_txt is not None:
                    FIGURE_COUNTER[0] += 1
                    cp = doc.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    try:
                        cp.style = "Caption"
                    except Exception:
                        pass
                    fig_clean = strip_fmt(fig_cap_txt).replace('"', "'")
                    fig_caption = f"{CONFIG.figure_name} {FIGURE_COUNTER[0]}: " + fig_clean
                    _run(cp, fig_caption,
                         italic=True, size=PT_SMALL, color=C_BLACK)
                    # Hidden TC field for List of Figures
                    tc_text = f"{CONFIG.figure_name} {FIGURE_COUNTER[0]}\t{fig_clean}"
                    _add_tc_field(cp, tc_text, "F")

            elif env_name in ("equation", "equation*", "align", "align*", "gather", "gather*", "multline", "multline*", "eqnarray", "eqnarray*"):
                full_env = f"\\begin{{{env_name}}}" + inner + f"\\end{{{env_name}}}"
                # Pandoc cannot convert display math containing \label{...};
                # strip labels before sending to Pandoc.
                full_env_for_pandoc = re.sub(r"\\label\{[^}]*\}", "", full_env)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                omml = None
                if TEMP_DIR and PANDOC_PATH:
                    try:
                        omml = latex_to_omml_element(full_env_for_pandoc, TEMP_DIR, PANDOC_PATH)
                    except Exception:
                        pass
                if omml is not None:
                    p._p.append(omml)
                else:
                    _run(p, strip_fmt(inner), italic=True, size=PT_NORM)

            elif env_name == "tcolorbox":
                # title from options [title={...}]
                opt_str = text[pos + env_m.end():]
                opt_m   = re.match(r"\[([^\]]*)\]", opt_str)
                box_title = ""
                if opt_m:
                    tm = re.search(r"title=\{?([^,}\]]+)\}?", opt_m.group(1))
                    if tm:
                        box_title = strip_fmt(tm.group(1))
                if box_title:
                    p = doc.add_paragraph()
                    _run(p, box_title, bold=True, size=PT_NORM)
                _walk(doc, inner, base_dir)

            elif env_name == "minipage":
                # extract_env leaves the mandatory {width} argument in inner;
                # strip it and render the real content.
                inner = inner.lstrip()
                if inner.startswith("{"):
                    _, inner = _strip_column_spec(inner)
                _walk(doc, inner, base_dir)

            elif env_name in ("document",):
                _walk(doc, inner, base_dir)
                end_pos = len(text)   # consumed everything

            # else: unknown env — walk its inner content
            else:
                _walk(doc, inner, base_dir)

            pos = end_pos
            continue

        # \includegraphics outside a figure environment (e.g. in center/minipage)
        img_m = re.match(r"\\includegraphics(?:\[([^\]]*)\])?\{([^}]*)\}", text[pos:])
        if img_m:
            flush()
            opts = img_m.group(1) or ""
            img_name = img_m.group(2)
            img_path = resolve_image_path(img_name, base_dir)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if img_path and img_path.exists():
                try:
                    from PIL import Image
                    img = Image.open(str(img_path))
                    px_w, px_h = img.size
                    nat_w = px_w / 96.0
                    nat_h = px_h / 96.0
                    w, h = _parse_graphics_size(opts, img_path)
                    if h is not None and w is None:
                        w = nat_w * (h / nat_h)
                    if w is None:
                        w = 4.0
                        h = nat_h * (w / nat_w)
                    if w > TEXT_WIDTH_INCHES:
                        scale = TEXT_WIDTH_INCHES / w
                        w *= scale
                        h *= scale
                    p.add_run().add_picture(str(img_path), width=Inches(w), height=Inches(h))
                except Exception:
                    _run(p, f"[Figure: {img_name}]", italic=True)
            else:
                _run(p, f"[Figure: {img_name}]", italic=True)
            pos += img_m.end()
            continue

        # \end{...} — should not appear here if extract_env worked, but skip it
        end_m = re.match(r"\\end\{[^}]*\}", text[pos:])
        if end_m:
            pos += end_m.end()
            continue

        # Double blank line → paragraph break
        blank = re.match(r"\n[ \t]*\n", text[pos:])
        if blank:
            flush()
            pos += blank.end()
            continue

        # Advance to next structural boundary or gather text chunk
        nxt = _STRUCT.search(text, pos)
        # Also stop at display math and double blanks
        dm2  = _DISP_MATH.search(text, pos)
        nxt2 = re.search(r"\n[ \t]*\n", text[pos:])

        candidates = [c for c in [
            nxt.start()          if nxt  else None,
            dm2.start()          if dm2  else None,
            pos + nxt2.start()   if nxt2 else None,
        ] if c is not None]

        stop = min(candidates) if candidates else n
        if stop == pos:
            # Avoid infinite loop
            pos += 1
            continue

        chunk = text[pos:stop]
        # Convert LaTeX line breaks \\[...] to paragraph breaks
        chunk = re.sub(r"\\\\(?:\[[^\]]*\])?", "\n\n", chunk)
        # Convert LaTeX newlines to paragraph breaks
        chunk = chunk.replace("\\newline", "\n\n")
        pos   = stop

        # Collect non-empty lines
        _VAR_LINE = re.compile(r"^(\$[^$]+\$|\\(?:gls|Gls)\{[^}]*\})")
        prev_was_var = False
        for line in chunk.split("\n"):
            line = line.strip()
            if not line:
                flush()
                prev_was_var = False
                continue
            # skip pure layout/control commands
            if re.match(
                r"\\(?:maketitle|tableofcontents|vspace|hspace|noindent|centering"
                r"|par\b|medskip|bigskip|smallskip|clearpage|newpage"
                r"|label|ref|pageref|addcontentsline|setlength"
                r"|pagestyle|thispagestyle|fancyhf|lhead|rhead|cfoot|headheight"
                r"|usepackage|documentclass|title|author|date"
                r"|usetikzlibrary|tcbuselibrary|hypersetup|rowcolors"
                r"|arrayrulecolor|cellcolor|newcounter|setcounter"
                r"|newcommand|renewcommand|setlist)\b",
                line
            ):
                continue
            is_var = bool(_VAR_LINE.match(line))
            if is_var and prev_was_var:
                flush()
            pending.append(line)
            prev_was_var = is_var

    flush()


def _maybe_add_tab_stop(para, text: str):
    """Add a tab stop if the cleaned text contains tab characters."""
    if "\t" in _clean(text):
        para.paragraph_format.tab_stops.add_tab_stop(Inches(1.5))


def _parse_leading_switches(text: str, bold: bool, italic: bool, size: int):
    r"""Extract initial font switches (\bfseries, \large, etc.) and update state."""
    size_map = {
        "tiny": 6, "scriptsize": 8, "footnotesize": 9, "small": 10,
        "normalsize": PT_NORM, "large": 14, "Large": 16, "LARGE": 20,
        "huge": 24, "Huge": 28,
    }
    pattern = re.compile(
        r"^\s*\\(bfseries|normalfont|itshape|"
        r"tiny|scriptsize|footnotesize|small|normalsize|large|Large|LARGE|huge|Huge)\b\s*"
    )
    while True:
        m = pattern.match(text)
        if not m:
            break
        cmd = m.group(1)
        if cmd == "bfseries":
            bold = True
        elif cmd == "normalfont":
            bold = False
            italic = False
        elif cmd == "itshape":
            italic = True
        elif cmd in size_map:
            size = size_map[cmd]
        text = text[m.end():]
    return bold, italic, size, text


def _scan_format_switches(text: str, bold: bool, italic: bool, size: int):
    r"""Scan text for font switches and return the final formatting state."""
    size_map = {
        "tiny": 6, "scriptsize": 8, "footnotesize": 9, "small": 10,
        "normalsize": PT_NORM, "large": 14, "Large": 16, "LARGE": 20,
        "huge": 24, "Huge": 28,
    }
    pattern = re.compile(
        r"\\(bfseries|normalfont|itshape|"
        r"tiny|scriptsize|footnotesize|small|normalsize|large|Large|LARGE|huge|Huge)\b"
    )
    for m in pattern.finditer(text):
        cmd = m.group(1)
        if cmd == "bfseries":
            bold = True
        elif cmd == "normalfont":
            bold = False
            italic = False
        elif cmd == "itshape":
            italic = True
        elif cmd in size_map:
            size = size_map[cmd]
    return bold, italic, size


def _add_para(doc: Document, raw: str, init_bold=False, init_italic=False, init_size=None):
    r"""Create justified paragraph(s) from raw inline LaTeX.

    Handles paragraph breaks (\\ or blank lines) and propagates font switches
    such as \bfseries / \large across the generated paragraphs.
    """
    raw = raw.strip()
    if not raw:
        return None
    paragraphs = re.split(r"\n\s*\n", raw)
    cur_bold = bool(init_bold)
    cur_italic = bool(init_italic)
    cur_size = init_size if init_size is not None else PT_NORM
    last_p = None
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        cur_bold, cur_italic, cur_size, para_text = _parse_leading_switches(
            para_text, cur_bold, cur_italic, cur_size
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        parse_inline(p, para_text, base_sz=PT_NORM,
                     bold=cur_bold, italic=cur_italic, size=cur_size)
        _maybe_add_tab_stop(p, para_text)
        last_p = p
    return last_p


# ─────────────────────────────────────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────────────────────────────────────

def extract_tables_to_temp(source: str, temp_dir: Path) -> int:
    """Extrae tablas del LaTeX en orden de aparicion y las guarda en carpeta temporal.

    Debe coincidir exactamente con lo que `_walk` renderiza como tabla:
      - entornos table/table* que contienen tabular/tabularx
      - entornos table/table* que contienen $\\begin{array}...\\end{array}$
      - entornos longtable
      - entornos tabular/tabularx directos (sin wrapper table)
    """
    table_count = 0

    # Fase 1: entornos table/table*/longtable
    wrapper_ranges = []   # (start, end)
    wrapper_tables = []   # (start, env_name, content)

    for base in ("table", "longtable"):
        pos = 0
        while True:
            idx = source.find(f"\\begin{{{base}}}", pos)
            if idx == -1:
                break
            # Permite table* / longtable* gracias a extract_env(env_base, ...)
            m = re.match(rf"\\begin\{{{re.escape(base)}\*?\}}", source[idx:])
            if not m:
                pos = idx + 1
                continue
            result = extract_env(source, base, idx)
            if not result:
                pos = idx + 1
                continue
            start, end, inner = result
            wrapper_ranges.append((start, end))

            if base == "longtable":
                wrapper_tables.append((start, "longtable", inner))
            else:
                # table/table*: preferir tabular/tabularx interno
                found = False
                for tenv in ("tabularx", "tabular"):
                    tr = extract_env(inner, tenv)
                    if tr:
                        _, _, tab_inner = tr
                        wrapper_tables.append((start, tenv, tab_inner))
                        found = True
                        break
                if not found:
                    # Tabla definida como \begin{array} dentro de math display
                    arr_match = re.search(
                        r"\$\s*\\begin\{array\}(.*?)\\end\{array\}\s*\$",
                        inner, re.DOTALL
                    )
                    if arr_match:
                        wrapper_tables.append((start, "array", arr_match.group(1)))
            pos = end

    # Fase 2: tabular/tabularx directos (no dentro de table/longtable)
    direct_tables = []
    for tenv in ("tabularx", "tabular"):
        pos = 0
        while True:
            idx = source.find(f"\\begin{{{tenv}}}", pos)
            if idx == -1:
                break
            result = extract_env(source, tenv, idx)
            if not result:
                pos = idx + 1
                continue
            start, end, inner = result
            inside_wrapper = any(ws <= start < we for ws, we in wrapper_ranges)
            if not inside_wrapper:
                direct_tables.append((start, tenv, inner))
            pos = end

    # Ordenar todo por posición en el source
    all_tables = wrapper_tables + direct_tables
    all_tables.sort(key=lambda x: x[0])

    for _, env_name, content in all_tables:
        table_count += 1

        # Remover \resizebox si existe (como en la tabla de presupuesto)
        # \resizebox{width}{height}{content} -> solo content
        content = re.sub(r"\\resizebox\{[^}]*\}\{[^}]*\}\{?", "", content)
        content = content.rstrip("}")  # Remover el cierre del resizebox si existe

        # Normalizar el especificador de columnas: Pandoc no entiende bien @{} y
        # genera una sola fila con todas las celdas. Se eliminan esos separadores
        # antes de delegar la conversión a Pandoc.
        col_spec, rest = _strip_column_spec(content)
        if col_spec:
            clean_spec = re.sub(r"@\{[^}]*\}", "", col_spec)
            content = "{" + clean_spec + "}" + rest

        # Aplanar tabulares anidados de una sola celda (ej. encabezados centrados
        # con \begin{tabular}[c]{@{}c@{}}texto\end{tabular}).
        content = _flatten_single_cell_tabulars(content)

        # Los entornos array se renderizan manualmente en _walk (is_array=True),
        # pero necesitamos un archivo Pandoc "placeholder" para mantener el
        # TABLE_COUNTER sincronizado. Convertimos el array a tabular para Pandoc.
        write_env = env_name
        if env_name == "array":
            write_env = "tabular"

        tex_file = temp_dir / f"tabla_{table_count:02d}_{env_name}.tex"
        tex_file.write_text(
            f"\\begin{{{write_env}}}" + content + f"\\end{{{write_env}}}",
            encoding="utf-8"
        )
    return table_count


def convert_tables_with_pandoc(temp_dir: Path, pandoc_path: str = "pandoc") -> bool:
    """Convierte todas las tablas .tex en temp_dir a .docx con Pandoc.
    Reintenta una vez si Pandoc no genera el archivo de salida."""
    tex_files = list(temp_dir.glob("tabla_*.tex"))
    if not tex_files:
        return False
    
    for tex_file in tex_files:
        # Nombre de salida: tabla_XX_pandoc.docx
        num_match = re.search(r'tabla_(\d+)_', tex_file.name)
        if not num_match:
            continue
        num = num_match.group(1)
        out_file = temp_dir / f"tabla_{num}_pandoc.docx"
        cmd = [pandoc_path, "-f", "latex", "-t", "docx", str(tex_file), "-o", str(out_file)]
        for attempt in range(2):
            try:
                result = subprocess.run(cmd, capture_output=True, text=True,
                                        check=False, timeout=60)
                if result.returncode == 0 and out_file.exists():
                    break
                if attempt == 0:
                    continue
                print(f"  [WARN] Pandoc falló para {tex_file.name}: {result.stderr.strip()}")
            except Exception as e:
                if attempt == 0:
                    continue
                print(f"  [WARN] Error ejecutando Pandoc para {tex_file.name}: {e}")
    return True


def expand_inputs(source: str, base_dir: Path, root_dir: Path = None, depth: int = 0, max_depth: int = 10) -> str:
    r"""Expand \input{file} directives by inlining the referenced files.
    Also handles \IfFileExists{file}{true}{false} by evaluating file existence
    against both the local directory and the project's root directory.
    """
    if depth > max_depth:
        return source
    
    root_dir = root_dir or base_dir
    
    def _resolve_file(filename: str):
        """Find file and return (content, file_dir) or (None, None)."""
        for d in (base_dir, root_dir):
            for ext in ("", ".tex"):
                cand = d / (filename + ext)
                if cand.exists():
                    return cand.read_text(encoding="utf-8"), cand.parent
        return None, None
    
    # 1. Handle \IfFileExists{file}{true}{false}
    def replace_iff(m):
        filename = m.group(1).strip()
        true_br = m.group(2)
        false_br = m.group(3) if m.group(3) is not None else ""
        content, _ = _resolve_file(filename)
        if content is not None:
            return true_br
        return false_br
    
    source = re.sub(
        r"\\IfFileExists\s*\{([^}]+)\}\s*\{((?:[^{}]|\{[^{}]*\})*)\}\s*\{((?:[^{}]|\{[^{}]*\})*)\}",
        replace_iff, source
    )
    
    # 2. Handle \input{file}
    def replace_input(m):
        filename = m.group(1).strip()
        content, file_dir = _resolve_file(filename)
        if content is not None:
            return expand_inputs(content, file_dir, root_dir, depth + 1, max_depth)
        return m.group(0)
    
    return re.sub(r"\\input\{([^}]+)\}", replace_input, source)


def main():
    global TEMP_DIR, PANDOC_PATH, CONFIG, PREAMBLE_GLOBAL

    in_path  = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("linea_base_en.tex")
    if len(sys.argv) > 2:
        out_path = Path(sys.argv[2])
    else:
        out_path = in_path.with_suffix(".docx")
    # If output path is relative, save next to the input .tex
    if not out_path.is_absolute():
        out_path = in_path.parent / out_path

    if not in_path.exists():
        print(f"Error: file not found: {in_path}", file=sys.stderr)
        sys.exit(1)

    # Crear carpeta temporal para archivos intermedios
    temp_dir = tempfile.mkdtemp(prefix="latex_conv_")
    TEMP_DIR = Path(temp_dir)

    # Fresh configuration for this run
    CONFIG = Config()

    try:
        raw_source = in_path.read_text(encoding="utf-8")
        # Strip comments BEFORE expanding inputs, so commented-out \input / \IfFileExists are ignored
        source = re.sub(r"(?m)(?<!\\)%.*$", "", raw_source)
        source = expand_inputs(source, in_path.parent, root_dir=in_path.parent)
        print(f"[INFO] Inputs expandidos: {len(source)} caracteres")

        # 1. Extraer tablas a carpeta temporal
        print("[1/4] Extrayendo tablas...")
        n_tables = extract_tables_to_temp(source, TEMP_DIR)
        print(f"       {n_tables} tablas encontradas")

        # 2. Parse preamble → populate MACROS and CONFIG
        print("[2/4] Leyendo preámbulo...")
        preamble_end = source.find("\\begin{document}")
        preamble = source[:preamble_end] if preamble_end != -1 else source
        PREAMBLE_GLOBAL = preamble
        parse_preamble(preamble, CONFIG)
        parse_geometry(preamble, CONFIG)

        # 3. Find external tools
        CONFIG.pandoc_path = _find_executable("pandoc", [
            r"C:\Program Files\Pandoc\pandoc.exe",
            r"C:\ProgramData\chocolatey\bin\pandoc.exe",
        ])
        CONFIG.pdflatex_path = _find_executable("pdflatex", [
            r"C:\Program Files\MiKTeX\miktex\bin\x64\pdflatex.exe",
            r"C:\ProgramData\MiKTeX\miktex\bin\x64\pdflatex.exe",
            r"C:\Users\Alienware\AppData\Local\Programs\MiKTeX\miktex\bin\x64\pdflatex.exe",
        ])
        CONFIG.pdftoppm_path = _find_executable("pdftoppm", [
            r"C:\Program Files\MiKTeX\miktex\bin\x64\pdftoppm.exe",
            r"C:\ProgramData\MiKTeX\miktex\bin\x64\pdftoppm.exe",
            r"C:\Users\Alienware\AppData\Local\Programs\MiKTeX\miktex\bin\x64\pdftoppm.exe",
        ])
        PANDOC_PATH = CONFIG.pandoc_path

        # 4. Convertir tablas con Pandoc
        if n_tables > 0:
            print("[3/4] Convirtiendo tablas con Pandoc...")
            convert_tables_with_pandoc(TEMP_DIR, CONFIG.pandoc_path)

        # Extract header text from \rhead{...}
        header_text = ""
        # Find all \rhead occurrences and pick the first one with actual text
        for rhead_match in re.finditer(r"\\rhead\s*\{", preamble):
            start = rhead_match.end()
            # Find matching } (simple brace balance)
            depth = 1
            i = start
            while i < len(preamble) and depth > 0:
                if preamble[i] == '{':
                    depth += 1
                elif preamble[i] == '}':
                    depth -= 1
                i += 1
            if depth == 0:
                raw_header = preamble[start:i-1]
                raw_header = re.sub(r"\\small\s*", "", raw_header)
                raw_header = raw_header.replace('%', '').strip()
                # Remove \includegraphics commands (they are rendered as images, not text)
                raw_header = re.sub(r"\\includegraphics(?:\[[^\]]*\])?\{[^}]*\}", "", raw_header)
                candidate = strip_fmt(raw_header).strip()
                if candidate and not candidate.startswith('[') and 'icono' not in candidate.lower():
                    header_text = candidate
                    break

        # Clean header text: remove image path artifacts
        if header_text.strip().startswith('[') or 'icono' in header_text.lower():
            header_text = ""

        # 5. Build document
        print("[4/4] Generando documento Word...")
        doc = Document()
        setup_document(doc, CONFIG)

        # Buscar logo
        logo_path = _find_logo(in_path.parent, CONFIG)
        add_header(doc, logo_path or Path("__missing_logo__"), header_text)
        add_footer(doc)

        parse_body(doc, source, in_path.parent)

        # Clean up empty paragraphs left after tables
        _remove_empty_paragraphs_after_tables(doc)

        doc.save(str(out_path))
        print(f"[OK] Guardado: {out_path}")

    finally:
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == "__main__":
    main()