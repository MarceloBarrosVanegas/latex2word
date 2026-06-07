#!/usr/bin/env python3
"""
Post-procesa DOCX generado por Pandoc:
- Arregla header con logo OFC
- Arregla footer con número de página
- Aplica fuente y estilos consistentes
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Calibri"
PT_NORM = 11
PT_SMALL = 9

def _run(para, text: str, bold=False, italic=False, size=None, color=None):
    if not text:
        return None
    r = para.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size or PT_NORM)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r

def _add_page_number_field(run):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

def strip_fmt(text: str) -> str:
    """Remove LaTeX markup."""
    if not text:
        return ""
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\small\s+", "", text)
    text = text.replace("\\%", "%")
    text = text.replace("\\&", "&")
    text = text.replace("~", " ")
    return text.strip()

def add_header(doc, logo_path, right_text):
    """Add OFC header with logo and project name."""
    sec = doc.sections[0]
    header = sec.header
    header.is_linked_to_previous = False
    
    # Clear existing
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    
    # Table for logo + text
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    
    # Left: Logo
    left_cell = table.cell(0, 0)
    left_cell.width = Inches(1.5)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_para.paragraph_format.space_after = Pt(0.2)
    
    if logo_path and logo_path.exists():
        run_img = left_para.add_run()
        run_img.add_picture(str(logo_path), height=Cm(0.8))
    else:
        _run(left_para, "OCEANS", bold=True, size=PT_SMALL)
    
    # Right: Text
    right_cell = table.cell(0, 1)
    right_cell.width = Inches(5.0)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_after = Pt(0.2)
    _run(right_para, right_text, size=PT_SMALL)
    
    # Bottom border on cells
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), '808080')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

def add_footer(doc):
    """Add centered page number."""
    sec = doc.sections[0]
    footer = sec.footer
    footer.is_linked_to_previous = False
    
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    _add_page_number_field(run)
    run.font.name = FONT
    run.font.size = Pt(PT_SMALL)

def extract_header_text(tex_file):
    """Extract \\rhead content from LaTeX."""
    source = tex_file.read_text(encoding="utf-8")
    header_text = "Galapagos Water Project (Ecuador)"
    rhead_match = re.search(r"\\rhead\{%?\s*(.*?)\s*%?\}", source, re.DOTALL)
    if rhead_match:
        raw = rhead_match.group(1)
        raw = re.sub(r"\\small\s*", "", raw)
        raw = raw.replace('%', '').strip()
        header_text = strip_fmt(raw)
    return header_text

def fix_styles(doc):
    """Apply consistent font and justification."""
    # Normal style
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(PT_NORM)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Headings
    for name in ['Heading 1', 'Heading 2', 'Heading 3']:
        try:
            s = doc.styles[name]
            s.font.name = FONT
            s.font.color.rgb = RGBColor(0, 0, 0)
        except:
            pass

def main():
    if len(sys.argv) < 3:
        print("Usage: python fix_pandoc_output.py input.docx output.docx [source.tex]")
        sys.exit(1)
    
    in_docx = Path(sys.argv[1])
    out_docx = Path(sys.argv[2])
    tex_file = Path(sys.argv[3]) if len(sys.argv) > 3 else Path("linea_base_en.tex")
    
    if not in_docx.exists():
        print(f"Error: {in_docx} not found")
        sys.exit(1)
    
    # Load Pandoc output
    doc = Document(str(in_docx))
    
    # Extract header text from LaTeX
    header_text = "Galapagos Water Project (Ecuador)"
    if tex_file.exists():
        header_text = extract_header_text(tex_file)
    
    # Apply fixes
    logo_path = tex_file.parent / "latex_to_word_converter" / "images" / "logo.jpg"
    if not logo_path.exists():
        logo_path = tex_file.parent / "00_figs" / "logo.jpg"
    
    add_header(doc, logo_path, header_text)
    add_footer(doc)
    fix_styles(doc)
    
    doc.save(str(out_docx))
    print(f"[OK] Fixed: {out_docx}")

if __name__ == "__main__":
    main()
