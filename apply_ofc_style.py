#!/usr/bin/env python3
"""
Aplica estilos OFC al documento Pandoc:
- Fuente Calibri en TODO
- Header con logo
- Footer con número
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

def apply_styles(doc):
    """Cambia TODO a Calibri."""
    # Estilos del documento
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        try:
            style = doc.styles[style_name]
            style.font.name = 'Calibri'
            if style_name == 'Normal':
                style.font.size = Pt(11)
            elif style_name == 'Heading 1':
                style.font.size = Pt(16)
                style.font.bold = True
            elif style_name == 'Heading 2':
                style.font.size = Pt(13)
                style.font.bold = True
            elif style_name == 'Heading 3':
                style.font.size = Pt(12)
                style.font.bold = True
        except:
            pass
    
    # Párrafos
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Calibri'
    
    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)

def add_header(doc, logo_path, header_text):
    """Header OFC con logo y texto."""
    sec = doc.sections[0]
    header = sec.header
    header.is_linked_to_previous = False
    
    # Limpiar
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    
    # Tabla logo | texto
    table = header.add_table(1, 2, width=Inches(6.5))
    table.autofit = False
    
    # Logo
    left = table.cell(0, 0)
    left.width = Inches(1.5)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    
    if logo_path.exists():
        run = p.add_run()
        run.add_picture(str(logo_path), height=Cm(0.8))
    else:
        run = p.add_run("OCEANS")
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.bold = True
    
    # Texto
    right = table.cell(0, 1)
    right.width = Inches(5.0)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(header_text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    
    # Línea
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), '808080')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)

def add_footer(doc):
    """Footer con número de página."""
    sec = doc.sections[0]
    footer = sec.footer
    footer.is_linked_to_previous = False
    
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    run.font.name = 'Calibri'
    run.font.size = Pt(9)

def main():
    # Cargar documento Pandoc
    doc = Document('temp_pandoc.docx')
    
    # Aplicar estilos OFC
    apply_styles(doc)
    
    # Header
    logo_path = Path('images/logo.jpg')
    if not logo_path.exists():
        logo_path = Path('../00_figs/logo.jpg')
    
    header_text = "Galapagos Water Project (Ecuador)"
    add_header(doc, logo_path, header_text)
    
    # Footer
    add_footer(doc)
    
    # Guardar
    doc.save('linea_base_en_WORD.docx')
    print('[OK] Estilos OFC aplicados a documento Pandoc')
    print('Archivo: linea_base_en_WORD.docx')

if __name__ == '__main__':
    main()
