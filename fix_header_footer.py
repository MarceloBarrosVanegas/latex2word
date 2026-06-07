#!/usr/bin/env python3
"""
Post-procesa DOCX de Pandoc:
- Agrega header OFC con logo
- Agrega footer con número de página
- Ajusta fuente a Calibri
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header_footer(doc, logo_path, header_text):
    """Agrega header OFC y footer con número de página."""
    
    # Configurar sección
    sec = doc.sections[0]
    
    # HEADER
    header = sec.header
    header.is_linked_to_previous = False
    
    # Limpiar header existente
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    
    # Tabla: Logo | Texto
    table = header.add_table(1, 2, width=Inches(6.5))
    table.autofit = False
    
    # Celda izquierda: Logo
    left = table.cell(0, 0)
    left.width = Inches(1.5)
    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.paragraph_format.space_after = Pt(2)
    
    if logo_path.exists():
        run = p_left.add_run()
        run.add_picture(str(logo_path), height=Cm(0.8))
    else:
        run = p_left.add_run("OCEANS")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.bold = True
    
    # Celda derecha: Texto
    right = table.cell(0, 1)
    right.width = Inches(5.0)
    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_after = Pt(2)
    run = p_right.add_run(header_text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    
    # Línea inferior
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['top', 'left', 'right']:
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), '808080')
        tcBorders.append(bottom)
        tcPr.append(tcBorders)
    
    # FOOTER - Número de página
    footer = sec.footer
    footer.is_linked_to_previous = False
    
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)
    
    fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    
    # Campo PAGE
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def fix_fonts(doc):
    """Cambia fuente a Calibri en todo el documento."""
    # Estilos
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3']:
        try:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            if style_name == 'Normal':
                style.font.size = Pt(11)
        except:
            pass
    
    # Párrafos existentes
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Calibri"
    
    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = "Calibri"


def main():
    if len(sys.argv) < 3:
        print("Uso: python fix_header_footer.py input.docx output.docx")
        return
    
    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2])
    tex_file = Path("linea_base_en.tex")
    
    if not input_file.exists():
        print(f"Error: No encuentro {input_file}")
        return
    
    # Cargar documento Pandoc
    doc = Document(str(input_file))
    
    # Extraer texto del header del LaTeX
    header_text = "Galapagos Water Project (Ecuador)"
    if tex_file.exists():
        content = tex_file.read_text(encoding="utf-8")
        import re
        m = re.search(r"\\rhead\{%?\s*\\small\s*(.+?)\s*%?\}", content, re.DOTALL)
        if m:
            header_text = m.group(1).replace("\\small", "").replace("%", "").strip()
    
    # Logo
    logo_path = Path("images/logo.jpg")
    if not logo_path.exists():
        logo_path = Path("../00_figs/logo.jpg")
    
    # Aplicar fixes
    add_header_footer(doc, logo_path, header_text)
    fix_fonts(doc)
    
    # Guardar
    doc.save(str(output_file))
    print(f"[OK] Guardado: {output_file}")


if __name__ == "__main__":
    main()
